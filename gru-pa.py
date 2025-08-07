# coding utf-8
import calendar
import datetime
import importlib.metadata
import os
import time

import nivo_chart as nc
import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
import streamlit.components.v1 as components
import streamlit_antd_components as sac
from chinese_calendar import is_workday
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from openpyxl.cell import MergedCell
from openpyxl.styles import Alignment, Border, Font, Side
from plotly.subplots import make_subplots
from pybadges import badge
from streamlit_condition_tree import condition_tree
from streamlit_extras.metric_cards import style_metric_cards
from streamlit_webrtc import WebRtcMode, webrtc_streamer
from wcwidth import wcswidth

from commFunc import (execute_sql, execute_sql_and_commit, get_update_content,
                      getUserEDKeys, getVerInfo, updatePyFileinfo)
from face_login import (clean_snapshot, face_login_cv, face_login_webrtc,
                        face_recognize_webrtc, update_face_data)
from gd_weather import get_city_weather
from gen_badges import gen_badge
from gen_license_plate import create_plate_image
from hf_weather import (get_city_aqi, get_city_history_weather,
                        get_city_now_weather, get_city_pf_weather,
                        get_city_warning_now)
from mysql_pool import get_connection

# cSpell:ignoreRegExp /[^\s]{16,}/
# cSpell:ignoreRegExp /\b[A-Z]{3,15}\b/g
# cSpell:ignoreRegExp /\b[A-Z]\b/g

@st.fragment
def login():
    st.set_page_config(layout="centered")
    face_type = None
    # 显示应用名称
    st.markdown(f"<font face='微软雅黑' color=purple size=20><center>**{APPNAME_CN}**</center></font>", unsafe_allow_html=True)
    # 登录表单容器
    login = st.empty()
    with login.container(border=True):
        userID, userCName = [], []
        sql = "SELECT DISTINCT(StationCN) from users order by StationCN"
        rows = execute_sql(cur, sql)
        station_type = st.selectbox(label="请选择站点", options=[row[0] for row in rows], index=0)
        sql = f"SELECT userID, userCName, StationCN from users where StationCN = '{station_type}' order by login_counter DESC, userCName"
        rows = execute_sql(cur, sql)
        for row in rows:
            userID.append(row[0])
            userCName.append(row[1])
        query_userCName = st.selectbox("请选择用户", userCName, index=None)
        st.session_state.password_login, login_index = True, 0
        if query_userCName is not None:
            userID = userID[userCName.index(query_userCName)]
            sql = f"SELECT param_value from users_setup where param_name = 'password_login' and userID = {userID}"
            row = execute_sql(cur, sql)
            if row:
                st.session_state.password_login = bool(row[0][0])
                if not st.session_state.password_login:
                    login_index = 1
        else:
            userID = None
        sql = f"SELECT Count(ID) from users_face_data where StationCN = '{station_type}'"
        cur.execute(sql)
        face_login_available = True if cur.fetchone()[0] > 0 else False
        # 用户密码输入框
        userPassword = st.text_input("请输入密码", max_chars=8, placeholder="用户初始密码为1234", type="password", autocomplete="off")
        login_type = sac.segmented(
            items=[
                sac.SegmentedItem(label="密码登录", disabled=not st.session_state.password_login),
                sac.SegmentedItem(label="人脸识别登录", disabled=not face_login_available),
            ], align="start", color='green', index=login_index
        )
        # 登录按钮
        buttonLogin = st.button("登录")

    # 如果点击了登录按钮
    result = None
    if buttonLogin:
        if login_type == "密码登录":
            # 如果用户编码和密码不为空
            if userID and userPassword:
                # 验证用户密码
                verifyUPW = verifyUserPW(userID, userPassword)
                # 如果密码验证成功
                if verifyUPW[0]:
                    #userPassword = verifyUPW[1]
                    sql = f"SELECT userID, userCName, userType, StationCN from users where userID = {userID} and userPassword = '{verifyUPW[1]}'"
                    result = execute_sql(cur, sql)
                elif not verifyUPW[0]:
                    st.error("登录失败, 请检查密码, 若忘记密码请联系管理员重置")
            else:
                st.warning("请选择用户并输入密码")
        elif login_type == "人脸识别登录":
            if st.session_state.client_local:
                st.info("正在启动人脸识别(local-cam), 请稍等...")
                face_type = 'cv'
                result = face_login_cv(station_type)
            else:
                face_type = 'web-cam'
                if face_type == 'web-cam':
                    st.info("正在启动人脸识别(web-cam), 请稍等...")
                    login.empty()
                    camera_capture(station_type)
                elif face_type == 'webrtc':
                    st.info("正在启动人脸识别(webrtc), 请稍等...")
                    face_type = 'webrtc'
                    st.session_state.login_webrtc = True
                    st.session_state.logged_in = True
                    st.session_state.StationCN = station_type
                    st.rerun()
                else:
                    st.error("人脸识别失败, 请使用密码登录")
    if result:
        login_init(result)
        login.empty()
        st.rerun()
    elif login_type == "人脸识别登录" and buttonLogin and face_type == 'cv':
        st.error("人脸识别失败, 请使用密码登录")


def login_init(result):
    st.empty()
    st.toast(f"用户: {result[0][1]} 登录成功, 欢迎回来")
    st.session_state.logged_in = True
    st.session_state.userID = result[0][0]
    st.session_state.userCName = result[0][1]
    st.session_state.userType = result[0][2]
    st.session_state.StationCN = result[0][3]
    st.session_state.userPwRechecked = False
    # 获取系统设置
    get_system_setup()
    # 获取城市编码
    get_city_code()
    # 自动获取历史天气, 免得过期后数据无法获取
    auto_get_history_weather()
    # 更新用户设置
    refresh_users_setup()
    # 更新系统访问次数
    sql = "UPDATE verinfo set pyLM = pyLM + 1 where pyFile = 'visitcounter'"
    execute_sql_and_commit(conn, cur, sql)
    updatePyFileinfo()
    # 更新版本信息
    verinfo, verLM = getVerInfo()
    app_version = f'{int(verinfo / 10000)}.{int((verinfo % 10000) / 100)}.{verinfo}'
    app_lm = time.strftime('%Y-%m-%d %H:%M', time.localtime(verLM))
    gen_badge(conn, cur, [], 'MySQL', APPNAME_EN, app_version, app_lm)
    # 更新用户访问次数
    sql = f"UPDATE users set login_counter = login_counter + 1 where userID = {st.session_state.userID}"
    execute_sql_and_commit(conn, cur, sql)
    now = datetime.datetime.now()
    valid_time = now.strftime("%Y-%m-%d")
    sql = f"SELECT notice from notices where StationCN = '{st.session_state.StationCN}' and start_time >= '{valid_time}' and '{valid_time}' <= end_time"
    result = execute_sql(cur, sql)
    # 获取限行信息
    if st.session_state.vehicle_restrict:
        st.session_state.vehicle_restrict_info = get_vehicle_restrict()
    else:
        st.session_state.vehicle_restrict_info = None
    if result or st.session_state.vehicle_restrict_info:
        st.session_state.menu_index = 0
    else:
        st.session_state.menu_index = 1
    # 删除超过MAXREVDAYS天的数据
    del_date = cal_date(-st.session_state.max_rev_days)
    sql = f"DELETE from pa_share where share_date <= '{del_date}'"
    execute_sql_and_commit(conn, cur, sql)
    st.set_page_config(layout="wide")


def logout():
    # 关闭游标
    cur.close()
    # 关闭数据库连接
    conn.close()

    # 清除会话状态中的所有键值对
    for key in st.session_state.keys():
        del st.session_state[key]

    # 重新运行当前脚本
    st.rerun()


def verifyUserPW(vUserName, vUserPW):
    st.session_state.userPwRechecked = False
    vUserEncPW = ""
    sql = f"SELECT userPassword from users where userID = {vUserName}"
    pwTable = execute_sql(cur, sql)
    if pwTable:
        vUserEncPW = pwTable[0][0]
        vUserDecPW = getUserEDKeys(vUserEncPW, "dec")
        if vUserPW == vUserDecPW:
            st.session_state.userPwRechecked = True

    return st.session_state.userPwRechecked, vUserEncPW


@st.fragment
def displayBigTimeCircle():
    components.html(open("./MyComponentsScript/Clock-Big-Circle.txt", "r", encoding="utf-8").read(), height=260)


@st.fragment
def displayVisitCounter():
    sql = "SELECT pyLM from verinfo where pyFile = 'visitcounter'"
    visitcount = execute_sql(cur, sql)[0][0]
    countScript = (open("./MyComponentsScript/FlipNumber.txt", "r", encoding="utf-8").read()).replace("visitcount", str(visitcount))
    components.html(countScript, height=40)


@st.fragment
def displaySmallTime():
    components.html(open("./MyComponentsScript/Clock-Small.txt", "r", encoding="utf-8").read(), height=34)


@st.fragment
def displaySmallClock():
    components.html(open("./MyComponentsScript/Clock-Number.txt", "r", encoding="utf-8").read(), height=30)


@st.fragment
def displayAppInfo(txt_height=300):
    infoStr = open("./MyComponentsScript/glowintext.txt", "r", encoding="utf-8").read()
    infoStr = infoStr.replace("软件名称", APPNAME_CN)
    verinfo, verLM = getVerInfo()
    infoStr = infoStr.replace("软件版本", f"软件版本: {int(verinfo / 10000)}.{int((verinfo % 10000) / 100)}.{verinfo} building {verinfo}")
    infoStr = infoStr.replace("更新时间", f"更新时间: {time.strftime('%Y-%m-%d %H:%M', time.localtime(verLM))}")
    update_type, update_content = get_update_content(f"./CHANGELOG.md")
    infoStr = infoStr.replace("更新内容", f"更新内容: {update_type} - {update_content}")
    components.html(infoStr, height=txt_height)


def changePassword():
    # 显示密码修改页面标题
    st.write("### :red[密码修改]")
    # 创建一个带有边框的容器
    changePW = st.empty()
    with changePW.container(border=True):
        # 输入原密码
        oldPassword = st.text_input("请输入原密码", max_chars=8, type="password", autocomplete="off")
        # 输入新密码
        newPassword = st.text_input("请输入新密码", max_chars=8, type="password", autocomplete="off")
        # 再次输入新密码以确认
        confirmPassword = st.text_input("请再次输入新密码", max_chars=8, placeholder="请确保和上一步输入的密码一致", type="password", autocomplete="new-password")
        # 确认修改按钮
        buttonSubmit = st.button("确认修改")

    # 检查原密码是否为空
    if oldPassword:
        # 验证用户原密码
        verifyUPW = verifyUserPW(st.session_state.userID, oldPassword)
        if verifyUPW[0]:
            oldPassword = verifyUPW[1]
        # 构造SQL查询语句，验证用户名和密码是否匹配
        sql = f"SELECT ID from users where userID = {st.session_state.userID} and userPassword = '{oldPassword}'"
        if execute_sql(cur, sql):
            # 检查新密码和确认密码是否填写且一致
            if newPassword and confirmPassword and newPassword != "":
                if newPassword == confirmPassword:
                    # 确认修改按钮是否被点击
                    if buttonSubmit:
                        # 加密新密码
                        newPassword = getUserEDKeys(newPassword, "enc")
                        # 构造SQL更新语句，更新用户密码
                        sql = f"UPDATE users set userPassword = '{newPassword}' where userID = {st.session_state.userID}"
                        # 执行SQL语句并提交
                        execute_sql_and_commit(conn, cur, sql)
                        # 显示密码修改成功提示，并要求重新登录
                        st.toast("密码修改成功, 请重新登录")
                        # 登出用户
                        logout()
                else:
                    # 显示密码不一致的错误信息
                    st.error("两次输入的密码不一致")
            else:
                # 显示新密码未填写的警告信息
                st.warning("请检查新密码")
        else:
            # 显示原密码错误的错误信息
            st.error("原密码不正确")
    else:
        st.warning("原密码不能为空")


@st.fragment
def changelog():
    changelogInfo = open("./CHANGELOG.md", "r", encoding="utf-8").read()
    st.markdown(changelogInfo)


def aboutReadme():
    new_content = ''
    package_pack = ['Streamlit', 'Streamlit-antd-components', 'NumPY', 'Pandas', 'Plotly', 'Python-docx', 'Openpyxl', 'XlsxWriter', 'PyJWT', 'Dlib', 'Face-recognition', 'Opencv-python', 'Streamlit-webrtc']
    with open('./README.md', 'r', encoding='utf-8') as file:
        lines = file.readlines()

    verinfo, verLM = getVerInfo()
    app_version = f'{int(verinfo / 10000)}.{int((verinfo % 10000) / 100)}.{verinfo}'
    app_lm = time.strftime('%Y-%m-%d %H:%M', time.localtime(verLM))
    for line in lines:
        if line.startswith("    ![GRU-PA ver]"):
            line = f"    ![GRU-PA ver](https://img.shields.io/badge/ver-{app_version}-blue.svg)"
        elif line.startswith("    ![GRU-PA updated]"):
            line = f"    ![GRU-PA updated](https://img.shields.io/badge/updated-{app_lm.replace('-', '/')[2:10]}%20{app_lm[-5:]}-orange.svg)"
        elif line.startswith("      !["):
            for each in package_pack:
                if line.startswith(f"      ![{each}]"):
                    package_ver = importlib.metadata.version(each)
                    line = f"      ![{each}](https://img.shields.io/badge/{each.replace('-', '_')}-{package_ver}-blue.svg)"
                    break
        new_content = new_content + line + "\n"
    new_content = new_content.replace('\n\n', '\n')
    with open("./README.md", "w", encoding='utf-8') as f:
        f.write(new_content)

    st.markdown(open("./README.md", "r", encoding="utf-8").read(), unsafe_allow_html=True)


def aboutInfo():
    updatePyFileinfo()
    cols_limit = 5
    st.subheader("关于本软件", divider="rainbow")
    st.subheader(":blue[Powered by Python and Streamlit]")
    module_pack = ['Python', 'MySQL', 'Streamlit', 'SAC', 'Pandas', 'NumPY', 'Plotly', 'Dlib', 'Openpyxl', 'Python-Docx']
    module_img = st.columns(cols_limit)
    for index, value in enumerate(module_pack):
        module_img[index % cols_limit].caption(value)
        module_img[index % cols_limit].image(f'./Images/logos/{value.replace(" ", "_")}.png')
    display_pypi()
    if st.context.theme.type == 'dark':
        st.write("###### :violet[为了获得更好的使用体验, 请使用浅色主题]")
    sac.divider(align="center", color="gray")
    st.image("./Images/badges/license-badge.svg")
    st.caption(":violet[Copyright © 2025 Simon. All rights reserved.]")
    st.image("./Images/logos/simon-logo.png", width=50)


def display_pypi():
    db_type = 'MySQL'
    cols_limit = 5
    pypi = st.columns(cols_limit)
    badge_pack = ['Streamlit', 'streamlit-antd-components', 'Pandas', 'NumPY', 'Plotly', 'Dlib', 'Openpyxl', 'Python-docx']
    verinfo, verLM = getVerInfo()
    app_version = f'{int(verinfo / 10000)}.{int((verinfo % 10000) / 100)}.{verinfo}'
    app_lm = time.strftime('%Y-%m-%d %H:%M', time.localtime(verLM))
    gen_badge(conn, cur, badge_pack, db_type, APPNAME_EN, app_version, app_lm)
    badge_pack = ['Python', db_type] + badge_pack
    for index, value in enumerate(badge_pack):
        pypi[index % cols_limit].image(f'./Images/badges/{value}-badge.svg')
    app_info_col = st.columns(cols_limit)
    app_info_col[0].image(f'./Images/badges/{APPNAME_EN}-badge.svg')
    app_info_col[1].image(f'./Images/badges/{APPNAME_EN}-lm-badge.svg')
    app_info_col[2].image('./Images/badges/build-badge.svg')


def get_md_task_status(task_date, userID, task_content):
    # 获取月份的第一天
    first_day_of_month = task_date.replace(day=1)
    # 使用 calendar.monthrange() 获取该月总天数，即最后一天
    _, last_day = calendar.monthrange(task_date.year, task_date.month)
    last_day_of_month = task_date.replace(day=last_day)

    sql = f"SELECT ID from clerk_work where clerk_work = '{task_content}' and task_date >= '{first_day_of_month}' and task_date <= '{last_day_of_month}' and clerk_id = {userID}"
    if not execute_sql(cur, sql):

        return True

    return False


@st.dialog("确认添加")
def confirm_add_task(task_date):
    st.markdown(f"### :red[请确认添加日期为: {task_date}]")
    btn_col = st.columns(2)
    btn_confirm = btn_col[0].button("确认", key="confirm_add_task", icon=":material/add_task:", use_container_width=True)
    btn_cancel = btn_col[1].button("取消", key="cancel_add_task", icon=":material/cancel:", use_container_width=True)
    if btn_confirm or btn_cancel:
        if btn_confirm:
            for key in st.session_state.keys():
                if key.startswith("task_work_") and st.session_state[key]:
                    #st.write(key, st.session_state[key])
                    temp_task_multi = 1
                    task_id = key[key.rfind("_") + 1:]
                    sql = f"SELECT pa_content, pa_score, task_group from gru_pa where ID = {task_id}"
                    task_result = execute_sql(cur, sql)
                    task_content, task_score, task_group = task_result[0]
                    if f'task_score_{task_id}' in st.session_state.keys():
                        task_score = st.session_state[f'task_score_{task_id}']
                    if f'task_multi_{task_id}' in st.session_state.keys():
                        task_score *= st.session_state[f'task_multi_{task_id}']
                        temp_task_multi = st.session_state[f'task_multi_{task_id}']
                    sql = f"SELECT ID from clerk_work where task_date = '{task_date}' and clerk_id = {st.session_state.userID} and clerk_work = '{task_content}' and task_group = '{task_group}'"
                    if not execute_sql(cur, sql):
                        sql = f"INSERT INTO clerk_work (task_date, clerk_id, clerk_cname, clerk_work, task_score, task_group, StationCN, task_multi) VALUES ('{task_date}', {st.session_state.userID}, '{st.session_state.userCName}', '{task_content}', {task_score}, '{task_group}', '{st.session_state.StationCN}', {temp_task_multi})"
                        execute_sql_and_commit(conn, cur, sql)
                        st.toast(f"工作量: [{task_content}] 分值: [{task_score}] 添加成功！")
                        sql = f"UPDATE pa_share set share_score = share_score - {task_score} where pa_ID = {task_id} and StationCN = '{st.session_state.StationCN}' and share_date = '{task_date}'"
                        execute_sql_and_commit(conn, cur, sql)
                    else:
                        st.toast(f"工作量: [{task_content}] 已存在！")
        st.rerun()


@st.fragment
def task_input():
    #st.markdown("### <font face='微软雅黑' color=red><center>工作量录入</center></font>", unsafe_allow_html=True)
    st.subheader("任务批量录入", divider="green")
    # 初始化任务组别图标
    task_group_icon = init_task_group_icon()
    # 刷新用户设置
    refresh_users_setup()
    # 更新用户工作组别频率
    update_users_group_frequency()
    col1, col2 = st.columns(2)
    col1.markdown(f"##### 当前用户: :blue[{st.session_state.userCName}]")
    with col1:
        flag_auto_task = sac.switch("自动选择日常工作", value=st.session_state.auto_task_check, align="start", on_label="On")
        flag_clerk_type = sac.switch("岗位工作类型", value=st.session_state.task_clerk_type, align="start", on_label="值班", off_label="白班")
    task_date = col2.date_input('工作时间', value=datetime.date.today() - datetime.timedelta(days=1), max_value="today")
    confirm_btn_input = st.button("确认添加")
    ttl_score = 0
    sql = f"SELECT clerk_work, task_score, task_group from clerk_work where clerk_id = {st.session_state.userID} and task_date = '{task_date}'"
    result = execute_sql(cur, sql)
    if result:
        st.markdown("##### 已输入工作量:\n\n")
        for row in result:
            st.write(f':violet[工作类型:] {row[2]} :orange[内容:] {row[0]} :green[分值:] {row[1]}')
            ttl_score += row[1]
        st.markdown(f':red[总分:] {ttl_score}')
    else:
        st.markdown(f'###### :red[无任何记录]')
    # 更新共享分
    update_pa_share(task_date)
    task_clerk_type = 1 if flag_clerk_type else 2
    expander_col = st.columns(2)
    # 常用任务
    with expander_col[0].expander(f"# :green[常用]", icon=':material/bookmark_star:', expanded=True):
        sql = f"SELECT ID, pa_content, pa_score, pa_group, multi_score, min_days, default_task, pa_share, task_type from gru_pa where task_valid = 1 and StationCN = '{st.session_state.StationCN}' and comm_task = {task_clerk_type} order by ID"
        rows2 = execute_sql(cur, sql)
        for row2 in rows2:
            show_task_list(row2, task_date, flag_auto_task, task_clerk_type)
    # 所有任务组别
    if st.session_state.task_group_sort:
        sql = f"""
            SELECT tg.task_group
            FROM (
                SELECT DISTINCT gru_pa.task_group
                FROM gru_pa
                WHERE gru_pa.StationCN = '{st.session_state.StationCN}'
            ) tg
            JOIN users_task_group_freq utgf ON tg.task_group = utgf.task_group
            WHERE utgf.userID = {st.session_state.userID}
            ORDER BY utgf.task_group_freq DESC
        """
    else:
        sql = f"SELECT DISTINCT(task_group) from gru_pa where StationCN = '{st.session_state.StationCN}'"
    rows = execute_sql(cur, sql)
    expander_col_index = 1
    for row in rows:
        sql = f"SELECT ID, pa_content, pa_score, pa_group, multi_score, min_days, default_task, pa_share, task_type from gru_pa where task_valid = 1 and StationCN = '{st.session_state.StationCN}' and task_group = '{row[0]}' and comm_task <> {task_clerk_type} order by ID"
        rows2 = execute_sql(cur, sql)
        if rows2:
            if row[0] in task_group_icon:
                expand_icon = task_group_icon[row[0]]
            else:
                expand_icon = 'dashboard_customize'
            with expander_col[expander_col_index % 2].expander(f"# :green[{row[0]}]", icon=f':material/{expand_icon}:', expanded=False):
                for row2 in rows2:
                    show_task_list(row2, task_date, flag_auto_task, task_clerk_type)
            expander_col_index += 1
    if confirm_btn_input:
        confirm_add_task(task_date)


def show_task_list(row2, task_date, flag_auto_task, task_clerk_type):
    if row2[5] == st.session_state.md_task_days:
        display_md_task = get_md_task_status(task_date, st.session_state.userID, row2[1])
    else:
        display_md_task = True
    auto_task = True if row2[6] == task_clerk_type and flag_auto_task else False
    if row2[7] == 0:
        if row2[4] == 1:
            title_score_info = ':blue[单次分值]'
        else:
            title_score_info = ':green[分值]'
    else:
        title_score_info = ':orange[总分值]'
    if row2[5] > 0 and display_md_task:
        st.checkbox(f":red[{row2[1]} {title_score_info}:{row2[2]}]", value=auto_task, key=f"task_work_{row2[0]}")
    elif display_md_task:
        if row2[1] == '每日记录检查':
            if st.session_state.userID in [1, 7, 11]:
                auto_task = True
        st.checkbox(f"{row2[1]} {title_score_info}:{row2[2]}", value=auto_task, key=f"task_work_{row2[0]}")
    task_col = st.columns(2)
    if row2[4] == 1:
        task_col[0].number_input(f":green[倍数]", min_value=1, max_value=10, value=1, step=1, key=f"task_multi_{row2[0]}")
    if row2[7] == 1:
        sql = f"SELECT clerk_cname, task_score from clerk_work where clerk_work = '{row2[1]}' and task_date = '{task_date}' and StationCN = '{st.session_state.StationCN}'"
        share_result = execute_sql(cur, sql)
        share_user_cname, share_user_score = [], []
        for share_row in share_result:
            share_user_cname.append(share_row[0])
            share_user_score.append(share_row[1])
        if share_user_cname:
            temp_share_info = ''
            for index, value in enumerate(share_user_cname):
                temp_share_info = temp_share_info + value + '' + str(share_user_score[index]) + '分/'
            task_col[0].markdown(f":violet[协作者: {temp_share_info[:-1]}]")
        sql = f"SELECT share_score from pa_share WHERE pa_id = {row2[0]} and share_date = '{task_date}'"
        cur.execute(sql)
        share_score = cur.fetchone()[0]
        if share_score < 0:
            share_score = 0
        if bool(row2[8]) and share_user_cname:
            share_score_value = share_score
        else:
            share_score_value = int(share_score / 2)
        if share_score == 0:
            task_col[0].markdown(f':blue[无剩余分值, 请与协作者协商更改]')
        task_col[0].number_input(label=":red[共享分值]", min_value=0, max_value=share_score, value=share_score_value, step=1, key=f"task_score_{row2[0]}", help=f"最大值{share_score}, 共享分值请与协作者协商后填写")


def update_pa_share(task_date):
    sql = f"INSERT INTO pa_share (pa_ID, pa_content, share_score, StationCN, share_date) SELECT ID, pa_content, pa_score, '{st.session_state.StationCN}', '{task_date}' from gru_pa where StationCN = '{st.session_state.StationCN}' and pa_share = 1 and pa_content not in (SELECT pa_content from pa_share where StationCN = '{st.session_state.StationCN}' and share_date = '{task_date}')"
    execute_sql_and_commit(conn, cur, sql)
    sql = f"SELECT pa_content from pa_share where StationCN = '{st.session_state.StationCN}' and share_date = '{task_date}'"
    rows = execute_sql(cur, sql)
    for row in rows:
        sql = f"SELECT pa_score from gru_pa where pa_content = '{row[0]}' and StationCN = '{st.session_state.StationCN}' and pa_share = 1"
        org_score = execute_sql(cur, sql)[0][0]
        sql = f"SELECT task_score from clerk_work where clerk_work = '{row[0]}' and StationCN = '{st.session_state.StationCN}' and task_date = '{task_date}'"
        results = execute_sql(cur, sql)
        for result in results:
            org_score = org_score - result[0]
        sql = f"UPDATE pa_share set share_score = {org_score} where pa_content = '{row[0]}' and StationCN = '{st.session_state.StationCN}' and share_date = '{task_date}'"
        execute_sql_and_commit(conn, cur, sql)


def query_task():
    #st.markdown("### <font face='微软雅黑' color=red><center>工作量查询及导出</center></font>", unsafe_allow_html=True)
    st.subheader("工作量查询及导出", divider="orange")
    col1, col2, col3 = st.columns(3)
    if st.session_state.userType == 'admin':
        userID, userCName = [], []
        sql = f"SELECT userID, userCName from users where StationCN = '{st.session_state.StationCN}' and clerk_pa <> 0 order by login_counter DESC, userCName"
        rows = execute_sql(cur, sql)
        for row in rows:
            userID.append(row[0])
            userCName.append(row[1])
        query_userCName = col1.selectbox("请选择查询用户", userCName)
        query_userID = userID[userCName.index(query_userCName)]
    elif st.session_state.userType == 'user':
        col1.markdown(f"##### 当前用户: :blue[{st.session_state.userCName}]")
        query_userCName = st.session_state.userCName
        query_userID = st.session_state.userID
    query_date_start = col2.date_input('查询开始时间', value=datetime.date.today() - datetime.timedelta(days=1), max_value="today")
    query_date_end = col3.date_input('查询结束时间', value=query_date_start, min_value=query_date_start, max_value="today")
    col4, col5, col6, col7 = st.columns(4)
    confirm_btn_output = col4.button("导出为Word文件")
    if st.session_state.userType == 'admin':
        confirm_btn_output_excel = col5.button("导出统计报表(全部人员)")
    else:
        confirm_btn_output_excel = False
    with col6:
        flag_approved = sac.switch("仅限已核定工作", value=True, on_label="On")
    with col7:
        flag_combine = sac.switch("合并统计", value=False, on_label="On")
    if flag_combine:
        sql_task = f"SELECT clerk_work, AVG(task_score) AS avg_task_score, task_group, count(clerk_work) FROM clerk_work WHERE task_approved >= {int(flag_approved)} and task_date >= '{query_date_start}' AND task_date <= '{query_date_end}' AND clerk_id = {query_userID} GROUP BY clerk_work, task_group ORDER BY task_group"
        affix_info, color_field, score_num = "(合并统计)", "单项分值", 1
    else:
        sql_task = f"SELECT task_date, clerk_work, task_score, task_group, task_approved from clerk_work where task_approved >= {int(flag_approved)} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' and clerk_id = {query_userID} order by task_date, task_group, ID, clerk_work"
        affix_info, color_field, score_num = "", "分值", 2
    display_area = st.empty()
    rows = execute_sql(cur, sql_task)
    if rows:
        ttl_score = 0
        for row in rows:
            if flag_combine:
                ttl_score = ttl_score + row[1] * row[3]
            else:
                ttl_score += row[2]
        ttl_score = int(ttl_score)
        df = pd.DataFrame(rows, dtype=str)
        if flag_combine:
            df.columns = ["工作", "单项分值", "工作组别", "工作项数"]
            for index, value in enumerate(rows):
                df.loc[index, "单项分值"] = int(float(df["单项分值"][index]))
        else:
            df.columns = ["日期", "工作", "分值", "工作组别", "核定状态"]
            for index, value in enumerate(rows):
                df.loc[index, "分值"] = int(df["分值"][index])
                df.loc[index, "核定状态"] = "已核定" if int(df["核定状态"][index]) == 1 else "未核定"
        with display_area.container():
            st.dataframe(df.style.apply(highlight_max, subset=[color_field]))
            st.markdown(f':green[共计:] {len(rows)}项工作{affix_info} :red[总分:] {ttl_score}')
    else:
        st.info(f":red[未查询到记录]")
    if flag_approved:
        approved_info = "全部已审核"
    else:
        approved_info = "包含未审核"
    if confirm_btn_output:
        headerFS = 14
        contentFS = 12
        quesDOC = Document()
        quesDOC.styles["Normal"].font.name = "Microsoft YaHei"
        quesDOC.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        pHeader = quesDOC.add_paragraph()
        pHeader.alignment = WD_ALIGN_PARAGRAPH.CENTER
        textHeader = pHeader.add_run(f"{query_userCName} {query_date_start} 至 {query_date_end} 工作量{affix_info}记录 {approved_info}", 0)
        textHeader.font.size = Pt(headerFS)
        textHeader.font.bold = True
        rows = execute_sql(cur, sql_task)
        if rows:
            i = 1
            for row in rows:
                pContent = quesDOC.add_paragraph()
                if flag_combine:
                    textContent = pContent.add_run(f"第{i}项 - 工作类型: {row[2]} 内容: {row[0]} 单项分值: {int(row[1])} 项数: {row[3]}次")
                else:
                    if row[4] == 1:
                        approved_text = "已核定"
                    else:
                        approved_text = "未核定"
                    textContent = pContent.add_run(f"第{i}项 - 日期: {row[0]} 工作类型: {row[3]} 内容: {row[1]} 分值: {row[2]} 状态: {approved_text}")
                textContent.font.size = Pt(contentFS)
                textContent.font.bold = False
                if row[score_num] < 0:
                    textContent.font.color.rgb = RGBColor(139, 0, 0)
                else:
                    textContent.font.color.rgb = RGBColor(0, 0, 0)
                i += 1
            sql = f"SELECT clerk_work, task_score, task_group from clerk_work where task_date >= '{query_date_start}' and task_date <= '{query_date_end}' and clerk_id = {query_userID} and task_score < 0 order by task_date, task_group, ID, clerk_work"
            deduct_result = execute_sql(cur, sql)
            if deduct_result:
                pContent = quesDOC.add_paragraph()
                textContent = pContent.add_run(f"其中共{len(deduct_result)}个减分项(未合并)")
                textContent.font.size = Pt(contentFS + 2)
                textContent.font.bold = True
                textContent.font.color.rgb = RGBColor(139, 0, 0)
                j = 1
                for deduct_row in deduct_result:
                    pContent = quesDOC.add_paragraph()
                    textContent = pContent.add_run(f"第{j}项 - 工作类型: {deduct_row[2]} 内容: {deduct_row[0]} 分值: {deduct_row[1]}")
                    textContent.font.size = Pt(contentFS)
                    textContent.font.bold = True
                    textContent.font.color.rgb = RGBColor(0, 0, 0)
                    j += 1
            for j in range(1):
                pContent = quesDOC.add_paragraph()
            pContent = quesDOC.add_paragraph()
            textContent = pContent.add_run(f"共计完成{i - 1}项工作{affix_info} 总分: {ttl_score}")
            textContent.font.size = Pt(contentFS + 2)
            textContent.font.bold = True
            textContent.font.color.rgb = RGBColor(155, 17, 30)
            if not flag_approved:
                for j in range(1):
                    pContent = quesDOC.add_paragraph()
                pContent = quesDOC.add_paragraph()
                textContent = pContent.add_run("核定签字:  ________")
            # 加入页码
            footer = quesDOC.sections[0].footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            add_page_number(paragraph)
            outputFile = f"./user_pa/{query_userCName}_{query_date_start}至{query_date_end}_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.docx"
            if os.path.exists(outputFile):
                os.remove(outputFile)
            quesDOC.save(outputFile)
            if os.path.exists(outputFile):
                if os.path.exists(outputFile):
                    with open(outputFile, "rb") as file:
                        content = file.read()
                    file.close()
                    buttonDL = st.download_button("点击下载", content, file_name=outputFile[outputFile.rfind("/") + 1:], icon=":material/download:", type="secondary")
                    st.success(f":green[成功导出至程序目录user_pa下 {outputFile[outputFile.rfind('/') + 1:]}]")
                    if buttonDL:
                        st.toast("文件已下载至你的默认目录")
            else:
                st.error(f":red[文件导出失败]")
        else:
            st.info(f":red[未查询到记录]")
    elif confirm_btn_output_excel:
        display_area.empty()
        sql = f"SELECT clerk_cname, clerk_work, AVG(task_score), COUNT(clerk_work), AVG(task_score) * COUNT(clerk_work), task_group FROM clerk_work WHERE task_approved >= {int(flag_approved)} AND task_date >= '{query_date_start}' AND task_date <= '{query_date_end}' and StationCN = '{st.session_state.StationCN}' GROUP BY clerk_cname, clerk_work, task_group ORDER BY clerk_cname"
        result = execute_sql(cur, sql)
        df = pd.DataFrame(result)
        df.columns = ["姓名", "工作项", "单项分值", "项数", "单项合计", "工作组"]
        for item in ['单项分值', '项数', '单项合计']:
            df[item] = pd.to_numeric(df[item], errors='coerce').astype(int)
        # 指定需要计算小计的列
        sum_columns = ["项数", "单项合计"]
        # 调用函数添加小计行
        df_with_subtotals = add_subtotals(df, "姓名", sum_columns)
        # 修正单项分值列
        df_with_subtotals['单项分值'] = pd.to_numeric(df_with_subtotals['单项分值'], errors='coerce')
        # 显示带有小计行的 DataFrame
        st.dataframe(df_with_subtotals)
        # 导出为 Excel
        outputFile = f"./user_pa/工作量统计表_{query_date_start}至{query_date_end}_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.xlsx"
        if os.path.exists(outputFile):
            os.remove(outputFile)
        with pd.ExcelWriter(outputFile, engine='openpyxl') as writer:
            df_with_subtotals.to_excel(writer, sheet_name='统计表', index=False, startrow=1)
            # 插入统计时间行
            report_date_range = f"工作量统计 时间：{query_date_start} 至 {query_date_end} {approved_info}"
            # 对excel文件进行格式化
            ws = writer.sheets['统计表']
            # 设置页面为横向
            ws.page_setup.orientation = 'landscape'
            # 添加页眉/页脚（页脚居中显示页码）
            ws.oddFooter.center.text = "&P / &N"
            # 冻结前两行为标题行
            ws.print_title_rows = '1:2'
            # 合并 A1:F1，并设置样式
            ws.merge_cells("A1:F1")
            cell = ws["A1"]
            cell.value = report_date_range
            cell.font = Font(name="微软雅黑", size=14, bold=True)
            # 创建一个通用的对齐设置
            alignment = Alignment(horizontal='center', vertical='center')
            # 设置标题行样式
            for cell in ws[1]:
                cell.font = Font(name="微软雅黑", size=12, bold=True)
            # 将“小计”行设为粗体，并设置字体大小
            for row in ws.iter_rows(min_row=2):  # 跳过标题行
                if isinstance(row[0].value, str) and row[0].value.startswith('小计'):
                    for cell in row:
                        cell.font = Font(name="微软雅黑", size=12, bold=True)
            # 设置正文其他行字体
            for row in ws.iter_rows(min_row=2):
                if not (isinstance(row[0].value, str) and row[0].value.startswith('小计')):
                    for cell in row:
                        cell.font = Font(name="微软雅黑", size=12)
            # 定义边框样式
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            special_columns = {
                "A": 15,
                "F": 22
            }
            for col in ws.columns:
                max_width = 0
                column = None
                for cell in col:
                    if isinstance(cell, MergedCell):
                        continue
                    column = cell.column_letter
                    value = str(cell.value) if cell.value else ""
                    # 使用 wcswidth 精确计算显示宽度
                    width = wcswidth(value)
                    if width > max_width:
                        max_width = width
                if column:
                    if column in special_columns:
                        ws.column_dimensions[column].width = special_columns[column]
                    else:
                        ws.column_dimensions[column].width = max_width + 2
            # 添加边框到所有单元格
            for row in ws.iter_rows():
                for cell in row:
                    cell.border = thin_border
                    if row[0].row == 1:
                        cell.alignment = alignment
                    else:
                        cell.alignment = Alignment(horizontal='left', vertical='center')

            # 创建只包含小计行的简报sheet
            df_brief = df_with_subtotals[df_with_subtotals['姓名'].str.startswith('小计')].copy()
            # 去掉姓名中的"小计: "前缀
            df_brief['姓名'] = df_brief['姓名'].str.replace('小计: ', '')
            # 重命名列以匹配要求：ID，姓名，工作项数，分值合计和备注
            df_brief.rename(columns={
                '姓名': '姓名',
                '项数': '工作项数',
                '单项合计': '分值合计'
            }, inplace=True)
            # 添加ID列和备注列
            df_brief.insert(0, 'ID', range(1, len(df_brief) + 1))
            df_brief['备注'] = ''
            # 只保留需要的列
            df_brief = df_brief[['ID', '姓名', '工作项数', '分值合计', '备注']]

            # 计算工作项数和分值合计的平均值
            avg_work_count = df_brief['工作项数'].mean()
            avg_score_total = df_brief['分值合计'].mean()

            # 添加平均值行
            avg_row = pd.DataFrame({
                'ID': [''],
                '姓名': ['平均值'],
                '工作项数': [avg_work_count],
                '分值合计': [avg_score_total],
                '备注': ['']
            })

            # 将平均值行添加到DataFrame
            df_brief = pd.concat([df_brief, avg_row], ignore_index=True)
            df_brief.to_excel(writer, sheet_name='简报', index=False, startrow=1)
            # 插入统计时间行
            report_date_range2 = f"工作量统计简报 {query_date_start} 至 {query_date_end}"
            # 对简报sheet进行格式化，使用与统计表相同的样式
            ws2 = writer.sheets['简报']
            # 合并 A1:F1，并设置样式
            ws2.merge_cells("A1:E1")
            cell = ws2["A1"]
            cell.value = report_date_range2
            # 设置页面为横向
            ws2.page_setup.orientation = 'landscape'
            # 添加页眉/页脚（页脚居中显示页码）
            ws2.oddFooter.center.text = "&P / &N"
            # 冻结第一行为标题行
            ws2.print_title_rows = '1:1'
            # 创建一个通用的对齐设置
            alignment = Alignment(horizontal='center', vertical='center')
            # 设置标题行样式
            for cell in ws2[1]:
                cell.font = Font(name="微软雅黑", size=14, bold=True)
            # 设置正文其他行字体
            for row in ws2.iter_rows(min_row=2):
                for cell in row:
                    cell.font = Font(name="微软雅黑", size=14)
            # 定义边框样式（使用与统计表相同的边框样式）
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            special_columns = {
                "A": 5,
                "B": 15,
                "C": 15,
                "D": 15,
                "E": 20
            }
            for col in ws2.columns:
                max_width = 0
                column = None
                for cell in col:
                    if isinstance(cell, MergedCell):
                        continue
                    column = cell.column_letter
                    value = str(cell.value) if cell.value else ""
                    # 使用 wcswidth 精确计算显示宽度
                    width = wcswidth(value)
                    if width > max_width:
                        max_width = width
                if column:
                    if column in special_columns:
                        ws2.column_dimensions[column].width = special_columns[column]
                    else:
                        ws2.column_dimensions[column].width = max_width + 2
            # 添加边框到所有单元格
            for row in ws2.iter_rows():
                for cell in row:
                    cell.border = thin_border
                    if row[0].row == 1:
                        cell.alignment = alignment
                    else:
                        cell.alignment = Alignment(horizontal='center', vertical='center')

        if os.path.exists(outputFile):
            with open(outputFile, "rb") as file:
                content = file.read()
            file.close()
            buttonDL = st.download_button("点击下载", content, file_name=outputFile[outputFile.rfind("/") + 1:], icon=":material/download:", type="secondary")
            st.success(f":green[成功导出至程序目录user_pa下 {outputFile[outputFile.rfind('/') + 1:]}]")
            if buttonDL:
                st.toast("文件已下载至你的默认目录")
        else:
            st.error(f":red[文件导出失败]")


def add_subtotals(df, group_column, sum_columns):
    # 创建一个空列表用于存储新的行数据
    new_rows = []
    # 遍历每个分组
    for name, group in df.groupby(group_column):
        # 计算当前分组的小计
        subtotal = group[sum_columns].sum()
        # 创建小计行
        subtotal_row = {col: "" for col in df.columns}
        subtotal_row[group_column] = f"小计: {name}"
        for col in sum_columns:
            subtotal_row[col] = subtotal[col]
        # 将小计行添加到列表中
        new_rows.append(subtotal_row)
    # 将小计行转换为 DataFrame
    subtotal_df = pd.DataFrame(new_rows)
    # 将原始 DataFrame 和 小计 DataFrame 合并
    result_df = pd.concat([df, subtotal_df], ignore_index=True)

    return result_df


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()

    # 插入 "第 "
    run.add_text("第 ")
    # 插入 PAGE 字段
    _new_page_field(run, "PAGE")
    run.add_text(" 页 / 共 ")
    # 插入 NUMPAGES 字段
    _new_page_field(run, "NUMPAGES")
    run.add_text(" 页")


def _new_page_field(run, field_code):
    """
    插入 PAGE 或 NUMPAGES 字段到指定 run 中
    :param run: 要插入的 run 对象
    :param field_code: 'PAGE' 或 'NUMPAGES'
    """
    # 创建 fldSimple 元素
    fld_elm = OxmlElement('w:fldSimple')
    # 设置属性
    fld_elm.set(qn('w:instr'), f'{field_code} \\* MERGEFORMAT')
    # 添加到 run
    run._r.append(fld_elm)


def manual_input():
    #st.markdown("### <font face='微软雅黑' color=red><center>工作量手工录入</center></font>", unsafe_allow_html=True)
    st.subheader("工作量手工录入", divider="green")
    items, clerk_type = [], ['默认', '值班', '白班']
    col1, col2, col3, col4 = st.columns(4)
    if st.session_state.userType == 'admin':
        userID, userCName = [], []
        sql = f"SELECT userID, userCName from users where StationCN = '{st.session_state.StationCN}' and clerk_pa <> 0 order by login_counter DESC, userCName"
        rows = execute_sql(cur, sql)
        for row in rows:
            userID.append(row[0])
            userCName.append(row[1])
        add_userCName = col1.selectbox("请选择查询用户", userCName)
        add_userID = userID[userCName.index(add_userCName)]
    elif st.session_state.userType == 'user':
        col1.markdown(f"##### 当前用户: :blue[{st.session_state.userCName}]")
        add_userCName = st.session_state.userCName
        add_userID = st.session_state.userID
    sql = f"SELECT DISTINCT(task_group) from gru_pa where StationCN = '{st.session_state.StationCN}'"
    rows = execute_sql(cur, sql)
    for row in rows:
        items.append(row[0])
    task_date = col2.date_input('工作时间', value=datetime.date.today() - datetime.timedelta(days=1), max_value="today")
    task_group = col3.selectbox('工作组别', items, index=None, accept_new_options=True)
    task_score = col4.number_input("单项分值", min_value=1, max_value=300, value=10, step=1)
    if st.session_state.userType == 'admin':
        with col1:
            flag_add_pa = sac.switch("加入固定列表", value=False, align="start", on_label="On")
            flag_default_task = sac.segmented(label="默认带入",
                items=clerk_type,
                align="start", size='sm', color='red', key='default_task', return_index=True
            )
        with col2:
            flag_multi_score = sac.switch("多倍计算", value=False, align="start", on_label="On")
            flag_comm_task = sac.segmented(label="常用工作",
                items=clerk_type,
                align="start", size='sm', color='red', key='comm_task', return_index=True
            )
        with col3:
            flag_share_score = sac.switch("共享分值", value=False, align="start", on_label="On")
        with col4:
            flag_task_type = sac.switch("共享独占", value=False, align="start", on_label="On")
        if flag_default_task > 0 and flag_default_task != flag_comm_task:
            flag_comm_task = flag_default_task
            st.markdown(f"##### :green[常用工作与默认带入值设置冲突, 已自动修正为{clerk_type[flag_comm_task]}]")
        if flag_multi_score and (flag_share_score or flag_task_type):
            flag_share_score, flag_task_type = False, False
            st.markdown(f"##### :green[多倍计算与共享分值、共享独占设置冲突, 已自动关闭共享分值和共享独占]")
        if flag_task_type and not flag_share_score:
            flag_share_score = True
            st.markdown(f"##### :green[共享独占与共享分值设置冲突, 已自动打开共享分值]")
    else:
        flag_add_pa, flag_multi_score, flag_comm_task, flag_default_task, flag_share_score, flag_task_type = False, False, False, False, False, False
    task_content = st.text_area("工作内容", height=100)
    confirm_btn_manual = st.button("确认添加")
    if task_group and task_content and confirm_btn_manual:
        sql = f"SELECT ID from clerk_work where task_date = '{task_date}' and clerk_id = {add_userID} and clerk_work = '{task_content}'"
        if not execute_sql(cur, sql):
            sql = f"INSERT INTO clerk_work (task_date, clerk_id, clerk_cname, clerk_work, task_score, task_group, StationCN) VALUES ('{task_date}', {add_userID}, '{add_userCName}', '{task_content}', {task_score}, '{task_group}', '{st.session_state.StationCN}')"
            execute_sql_and_commit(conn, cur, sql)
            st.toast(f"用户: :green[{add_userCName}] 工作量: :blue[{task_content}] 添加成功！")
            if flag_add_pa:
                sql = f"SELECT ID from gru_pa where StationCN = '{st.session_state.StationCN}' and pa_content = '{task_content}' and pa_score = {task_score}"
                if not execute_sql(cur, sql):
                    sql = f"INSERT INTO gru_pa (pa_content, pa_score, pa_group, task_group, multi_score, default_task, comm_task, StationCN, pa_share, task_type) VALUES ('{task_content}', {task_score}, '全员', '{task_group}', {int(flag_multi_score)}, {int(flag_default_task)}, {int(flag_comm_task)}, '{st.session_state.StationCN}', {int(flag_share_score)}, {int(flag_task_type)})"
                    execute_sql_and_commit(conn, cur, sql)
                    reset_table_num(True)
                    st.toast(f"工作量: :blue[{task_content}] 添加至列表成功！")
                else:
                    st.warning(f"工作量: :blue[{task_content}] 在列表中已存在！")
            if flag_share_score:
                update_pa_share(task_date)
        else:
            st.warning(f"工作量: :blue[{task_content}] 已存在！")
    elif not task_group:
        st.warning(f"请选择工作组！")
    elif not task_content:
        st.warning(f"请输入工作内容！")


def reset_table_num(flag_force=False):
    if not flag_force:
        confirm_btn_reset = st.button("确认重置PA-Number")
    else:
        confirm_btn_reset = True
    if confirm_btn_reset:
        for modify_table in ['gru_pa', 'gru_pa_deduct']:
            i, sql = 1, ''
            if modify_table == 'gru_pa':
                sql = f"SELECT ID from {modify_table} order by task_group, ID, pa_num"
            elif modify_table == 'gru_pa_deduct':
                sql = f"SELECT ID from {modify_table} order by ID"
            if sql:
                rows = execute_sql(cur, sql)
                for row in rows:
                    sql = f"UPDATE {modify_table} SET pa_num = {i} where ID = {row[0]}"
                    execute_sql_and_commit(conn, cur, sql)
                    i += 2
        if not flag_force:
            st.success("数据库PA-Number重置成功")


#@st.fragment
def task_modify():
    #st.markdown("### <font face='微软雅黑' color=red><center>记录修改</center></font>", unsafe_allow_html=True)
    st.subheader("记录修改", divider="red")
    if st.session_state.userType != 'admin':
        sac.alert("已核定的记录无法修改和删除, 如若修改请联系管理员", icon="warning", banner=sac.Banner(play=True, direction='left', speed=150, pauseOnHover=True), closable=True)
    col1, col2, col3, col4 = st.columns(4)
    if st.session_state.userType == 'admin':
        userID, userCName = [], []
        sql = f"SELECT userID, userCName from users where StationCN = '{st.session_state.StationCN}' and clerk_pa <> 0 order by login_counter DESC, userCName"
        rows = execute_sql(cur, sql)
        for row in rows:
            userID.append(row[0])
            userCName.append(row[1])
        query_userCName = col1.selectbox("请选择查询用户", userCName)
        query_userID = userID[userCName.index(query_userCName)]
    else:
        col1.markdown(f"##### 当前用户: :blue[{st.session_state.userCName}]")
        query_userID = st.session_state.userID
        query_userCName = st.session_state.userCName
    query_date_start = col2.date_input('查询开始时间', value=datetime.date.today() - datetime.timedelta(days=1), max_value="today")
    query_date_end = col3.date_input('查询结束时间', value=query_date_start, min_value=query_date_start, max_value="today")
    user_task_id_pack = []
    ttl_score = 0
    st.markdown(f'##### 已输入工作量:')
    # 已核定工作量
    sql = f"SELECT clerk_work, task_score, task_group, ID, task_approved from clerk_work where clerk_id = {query_userID} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' and task_approved = 1"
    result = execute_sql(cur, sql)
    for row in result:
        user_task_id_pack.append(row[3])
    with st.expander(":green[已核定]", icon=":material/checklist:", expanded=False):
        display_are = st.empty()
        with display_are.container():
            if result:
                for row in result:
                    st.write(f'ID:{row[3]} :violet[工作类型:] {row[2]} :orange[内容:] {row[0]} :green[分值:] {row[1]}')
                    ttl_score += row[1]
            else:
                st.markdown(f'###### :red[未查询到记录]')
    # 未核定工作量
    sql = f"SELECT clerk_work, task_score, task_group, ID, task_approved from clerk_work where clerk_id = {query_userID} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' and task_approved = 0"
    result = execute_sql(cur, sql)
    for row in result:
        user_task_id_pack.append(row[3])
    with st.expander(":blue[未核定]", icon=":material/notes:", expanded=True):
        display_are = st.empty()
        with display_are.container():
            if result:
                for row in result:
                    st.write(f'ID:{row[3]} :violet[工作类型:] {row[2]} :orange[内容:] {row[0]} :green[分值:] {row[1]}')
                    ttl_score += row[1]
            else:
                st.markdown(f'###### :red[未查询到记录]')
    st.markdown(f':red[总分:] {ttl_score}')
    task_modify_id = col4.selectbox("请选择任务ID", user_task_id_pack, index=None)
    if task_modify_id:
        confirm_btn_delete = col1.button("删除", type="primary")
        btn_modify = col3.button("修改", type="primary")
        sql = f"SELECT task_approved FROM clerk_work where ID = {task_modify_id} and clerk_id = {query_userID}"
        approved_result = bool(execute_sql(cur, sql)[0][0])
        if not approved_result or st.session_state.userType == 'admin':
            sql = f"SELECT clerk_work, task_score from clerk_work where ID = {task_modify_id} and clerk_id = {query_userID}"
            org_work, org_score = execute_sql(cur, sql)[0]
            form = st.columns(3)
            pa_share_results = get_rem_share_score(task_modify_id, query_userID)
            if pa_share_results[0]:
                modify_min_value = 0
                modify_max_value = pa_share_results[1] + pa_share_results[2]
            else:
                modify_min_value = st.session_state.max_deduct_score
                modify_max_value = 1000
            #st.write(org_score, modify_max_value, modify_min_value)
            modify_content = form[0].text_area("请输入修改后的内容", value=org_work, height=100)
            modify_score = form[1].number_input(f"请输入修改后的分值, 最大值{modify_max_value}", min_value=modify_min_value, max_value=modify_max_value, value=org_score, step=1)
            if confirm_btn_delete:
                col2.button("确认删除", type="secondary", on_click=delete_task, args=(task_modify_id, query_userID,))
            elif btn_modify and modify_content != '':
                col4.button("确认修改", type="secondary", on_click=modify_task, args=(task_modify_id, query_userID, modify_content, modify_score, pa_share_results))
            elif modify_content == '':
                st.info('请填写修改内容')
        else:
            st.error(f"ID:{task_modify_id} 被核定的记录无法修改, 请联系管理员!")
    else:
        st.info('请选择要处理的记录ID')


def delete_task(task_modify_id, query_userID):
    flag_pa_share, rem_share_score, org_score, org_date, pa_share_id, max_score = get_rem_share_score(task_modify_id, query_userID)
    if st.session_state.userType == 'admin':
        sql = f"DELETE FROM clerk_work where ID = {task_modify_id} and clerk_id = {query_userID}"
    else:
        sql = f"DELETE FROM clerk_work where ID = {task_modify_id} and clerk_id = {query_userID} and task_approved = 0"
    execute_sql_and_commit(conn, cur, sql)
    sql = f"SELECT ID FROM clerk_work where ID = {task_modify_id} and clerk_id = {query_userID}"
    if not execute_sql(cur, sql):
        if flag_pa_share:
            modify_share_score = rem_share_score + org_score
            if modify_share_score > max_score:
                sql = f"UPDATE pa_share set share_score = {max_score} where pa_ID = {pa_share_id} and share_date = '{org_date}' and StationCN = '{st.session_state.StationCN}'"
                st.toast(f"ID:{task_modify_id} 分值错误, 剩余共享分值不能大于共享总分{max_score}, 剩余共享分已恢复默认值!")
            else:
                sql = f"UPDATE pa_share set share_score = {modify_share_score} where pa_ID = {pa_share_id} and share_date = '{org_date}' and StationCN = '{st.session_state.StationCN}'"
                st.toast(f"ID:{task_modify_id} 删除成功!")
            execute_sql_and_commit(conn, cur, sql)
    else:
        st.toast(f"ID:{task_modify_id} 删除失败! 被核定的记录无法删除, 请联系管理员!")


def get_rem_share_score(task_modify_id, query_userID):
    sql = f"SELECT task_score, clerk_work, task_date from clerk_work where ID = {task_modify_id} and clerk_id = {query_userID}"
    org_score, org_work, org_date = execute_sql(cur, sql)[0]
    sql = f"SELECT pa_share, ID, pa_score from gru_pa where pa_content = '{org_work}' and StationCN = '{st.session_state.StationCN}'"
    pa_share_result = execute_sql(cur, sql)
    if pa_share_result:
        flag_pa_share, pa_share_id, max_score = pa_share_result[0]
        flag_pa_share = bool(flag_pa_share)
    else:
        flag_pa_share = False
    if flag_pa_share:
        sql = f"SELECT share_score from pa_share where pa_ID = {pa_share_id} and share_date = '{org_date}' and StationCN = '{st.session_state.StationCN}'"
        rem_pa_share_result = execute_sql(cur, sql)
        if rem_pa_share_result:
            rem_share_score = rem_pa_share_result[0][0]
            if rem_share_score < 0:
                rem_share_score = 0
            return True, rem_share_score, org_score, org_date, pa_share_id, max_score
        else:
            return False, None, None, None, None, None

    return False, None, None, None, None, None


def modify_task(task_modify_id, query_userID, modify_content, modify_score, flag_share_pack):
    flag_pa_share, rem_share_score, org_score, org_date, pa_share_id, max_score = flag_share_pack
    sql = f"UPDATE clerk_work SET clerk_work = '{modify_content}', task_score = {modify_score} where ID = {task_modify_id} and clerk_id = {query_userID}"
    execute_sql_and_commit(conn, cur, sql)
    if flag_pa_share:
        sql = f"UPDATE pa_share set share_score = share_score - {modify_score - org_score} where pa_ID = {pa_share_id} and share_date = '{org_date}' and StationCN = '{st.session_state.StationCN}'"
        execute_sql_and_commit(conn, cur, sql)
    sql = f"SELECT ID from clerk_work where clerk_work = '{modify_content}' and task_score = {modify_score} and ID = {task_modify_id} and clerk_id = {query_userID}"
    if execute_sql(cur, sql):
        st.toast(f"ID:{task_modify_id} 工作:{modify_content} 分值:{modify_score} 修改成功!")
    else:
        st.toast(f"ID:{task_modify_id} 修改失败! 请检查输入的内容及分值是否正确!")


@st.fragment
def check_data():
    #st.markdown("### <font face='微软雅黑' color=red><center>数据检查与核定</center></font>", unsafe_allow_html=True)
    st.subheader("数据检查与核定", divider="blue")
    col1, col2 = st.columns(2)
    userID, userCName = [], []
    sql = f"SELECT userID, userCName from users where StationCN = '{st.session_state.StationCN}' and clerk_pa <> 0 order by login_counter DESC, userCName"
    rows = execute_sql(cur, sql)
    for row in rows:
        userID.append(row[0])
        userCName.append(row[1])
    query_date_start = col1.date_input('查询开始时间', value=datetime.date.today() - datetime.timedelta(days=1), max_value="today")
    query_date_end = col2.date_input('查询结束时间', value=query_date_start, min_value=query_date_start, max_value="today")
    col = st.columns(4)
    dur_time = query_date_end - query_date_start
    st.markdown(f'##### 统计周期: {dur_time.days + 1}天')
    confirm_btn_approv = col[2].button("核定")
    confirm_btn_check = col[3].button("检查")
    cert_userCName = col[0].selectbox("请选择核定用户", userCName)
    cert_userID = userID[userCName.index(cert_userCName)]
    with col[1]:
        flag_all = sac.switch("全选", True)
    if confirm_btn_check:
        for index, value in enumerate(userID):
            sql = f"SELECT pa_content, min_days from gru_pa where StationCN = '{st.session_state.StationCN}' and min_days > 0 order by min_days DESC"
            rows = execute_sql(cur, sql)
            for row in rows:
                sql = f"SELECT count(ID) from clerk_work where clerk_work = '{row[0]}' and clerk_id = {value} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}'"
                task_count = execute_sql(cur, sql)[0][0]
                if task_count > 1 and task_count > dur_time.days / row[1]:
                    st.warning(f"用户: {userCName[index]} 工作: [{row[0]}] 应该 1次/{row[1]}天, 实际: {task_count}次 已超量, 请检查记录！")
    else:
        task_pack = []
        if flag_all:
            sql = f"SELECT ID, clerk_cname, task_date, clerk_work, task_score, task_group from clerk_work where StationCN = '{st.session_state.StationCN}' and task_approved = 0 and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' order by task_date, task_group, clerk_work, clerk_cname, task_score"
        else:
            sql = f"SELECT ID, clerk_cname, task_date, clerk_work, task_score, task_group from clerk_work where StationCN = '{st.session_state.StationCN}' and clerk_id = {cert_userID} and task_approved = 0 and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' order by task_date, task_group, clerk_work, clerk_cname, task_score"
        result = execute_sql(cur, sql)
        if result:
            for row in result:
                task_pack.append(f'{str(row[2])[5:]} {row[1]} {row[5]} 内容:{row[3]} 分值:{row[4]} ID:{row[0]}')
            approve_pack = sac.transfer(items=task_pack, label='工作量核定', titles=['项目'], reload=True, align='center', search=True, pagination=True, use_container_width=True)
            if confirm_btn_approv and approve_pack:
                for each in approve_pack:
                    approve_id = each[each.rfind('ID:') + 3:].strip()
                    sql = f"UPDATE clerk_work SET task_approved = 1 where ID = {approve_id}"
                    execute_sql_and_commit(conn, cur, sql)
                st.markdown(f"##### {len(approve_pack)} 个工作量已核定")
                sql = f"SELECT Count(ID) from clerk_work where task_approved = 1 and StationCN = '{st.session_state.StationCN}'"
                result = execute_sql(cur, sql)[0][0]
                sql = f"UPDATE verinfo set pyLM = {result} where pyFile = 'task_approved_counter'"
                execute_sql_and_commit(conn, cur, sql)
        else:
            st.markdown(f'###### :red[无任何记录]')


def resetPassword():
    # 显示副标题和分隔线
    st.subheader(":orange[密码重置]", divider="red")
    if st.session_state.userPwRechecked:
        # 显示重置用户信息提示
        st.write(":red[**重置用户信息**]")
        # 获取用户编码
        userID, userCName = [], []
        sql = f"SELECT userID, userCName from users where StationCN = '{st.session_state.StationCN}' order by login_counter DESC, userCName"
        rows = execute_sql(cur, sql)
        for row in rows:
            userID.append(row[0])
            userCName.append(row[1])
        query_userCName = st.selectbox("请选择用户", userCName, index=None)
        if query_userCName is not None:
            rUserName = userID[userCName.index(query_userCName)]
        if query_userCName is not None:
            # 创建重置按钮
            btnResetUserPW = st.button("重置", type="primary")
            if btnResetUserPW:
                st.button("确认重置", type="secondary", on_click=actionResetUserPW, args=(rUserName,))
                st.session_state.userPwRechecked = False
    else:
        vUserPW = st.text_input("请输入密码", max_chars=8, placeholder="请输入管理员密码, 以验证身份", type="password", autocomplete="off")
        # 检查是否输入了密码
        if vUserPW:
            # 验证密码
            if verifyUserPW(st.session_state.userID, vUserPW)[0]:
                st.session_state.userPwRechecked = True
                st.rerun()
            # 如果密码错误，显示错误提示
            else:
                st.error("密码错误, 请重新输入")
                st.session_state.userPwRechecked = False


def actionResetUserPW(rUserName):
    rInfo = ""

    # 获取用户加密密钥
    resetPW = getUserEDKeys("1234", "enc")
    # 构建 SQL 更新语句
    sql = f"UPDATE users SET userPassword = '{resetPW}' where userID = {rUserName}"
    # 执行 SQL 并提交
    execute_sql_and_commit(conn, cur, sql)
    # 更新信息，表示密码已重置
    rInfo += "密码已重置为: 1234"

    st.success(f"**{rInfo}**")


def deduction_input():
    #st.markdown("### <font face='微软雅黑' color=red><center>减分项录入</center></font>", unsafe_allow_html=True)
    st.subheader("减分项录入", divider="red")
    col1, col2 = st.columns(2)
    userID, userCName, pa_deduct, pa_deduct_score = [], [], [], []
    sql = f"SELECT userID, userCName from users where StationCN = '{st.session_state.StationCN}' and clerk_pa <> 0 order by login_counter DESC, userCName"
    rows = execute_sql(cur, sql)
    for row in rows:
        userID.append(row[0])
        userCName.append(row[1])
    deduct_userCName = col1.selectbox("请选择用户", userCName)
    deduct_userID = userID[userCName.index(deduct_userCName)]
    deduct_date = col2.date_input("请选择日期", datetime.date.today() - datetime.timedelta(days=1), max_value="today")
    sql = f"SELECT pa_content, pa_score from gru_pa_deduct where StationCN = '{st.session_state.StationCN}' order by pa_score, pa_content"
    rows = execute_sql(cur, sql)
    for row in rows:
        pa_deduct.append(row[0])
        pa_deduct_score.append(row[1])
    task_deduct = col1.selectbox("扣分项", pa_deduct, index=None)
    if task_deduct:
        task_score = pa_deduct_score[pa_deduct.index(task_deduct)]
    else:
        task_score = -50
    deduct_score = col2.number_input("扣分", min_value=st.session_state.max_deduct_score, max_value=-10, value=task_score, step=10)
    deduct_content = col1.text_area("自定义扣分项内容", value=task_deduct, placeholder="可选择固定扣分项后修改", height=100)
    confirm_btn_add = st.button("确认添加")
    if confirm_btn_add:
        #st.write(deduct_content, deduct_score, deduct_userID, deduct_userCName, deduct_date)
        if deduct_content:
            sql = f"INSERT INTO clerk_work (task_date, clerk_id, clerk_cname, clerk_work, task_score, task_group, task_approved, StationCN) VALUES ('{deduct_date}', {deduct_userID}, '{deduct_userCName}', '{deduct_content}', {deduct_score}, '扣分', 1, '{st.session_state.StationCN}')"
            execute_sql_and_commit(conn, cur, sql)
            st.success(f"用户: :blue[{deduct_userCName}] 扣分项: :red[{deduct_content}] 添加成功")
            sql = f"SELECT ID from gru_pa_deduct where pa_content = '{deduct_content}' and pa_score = {deduct_score} and StationCN = '{st.session_state.StationCN}'"
            if not execute_sql(cur, sql):
                sql = f"INSERT INTO gru_pa_deduct(pa_content, pa_score, StationCN) VALUES ('{deduct_content}', {deduct_score}, '{st.session_state.StationCN}')"
                execute_sql_and_commit(conn, cur, sql)
                st.success(f"扣分项: :red[{deduct_content}] 已添加至固定列表")
                reset_table_num(True)
        else:
            st.error("请输入扣分项内容")


def highlight_max(x, forecolor='black', backcolor="#D61919"):
    is_max = x == x.max()

    return [f'color: {forecolor}; background-color: {backcolor}' if v else '' for v in is_max]


def gen_chart():
    #st.markdown("### <font face='微软雅黑' color=red><center>趋势图</center></font>", unsafe_allow_html=True)
    st.subheader("趋势图", divider="rainbow")
    col1, col2, col3 = st.columns(3)
    tab1, tab2 = st.tabs(["📊 图表", "🧮 数据"])
    #tab1, tab2 = st.tabs(["📈 图表", "📋 数据"])
    if st.session_state.userType == 'admin':
        userID, userCName = [], []
        sql = f"SELECT userID, userCName from users where StationCN = '{st.session_state.StationCN}' and clerk_pa <> 0 order by login_counter DESC, userCName"
        rows = execute_sql(cur, sql)
        for row in rows:
            userID.append(row[0])
            userCName.append(row[1])
        query_userCName = col1.selectbox("请选择查询用户", userCName)
        query_userID = userID[userCName.index(query_userCName)]
    elif st.session_state.userType == 'user':
        col1.markdown(f"##### 当前用户: :blue[{st.session_state.userCName}]")
        query_userCName = st.session_state.userCName
        query_userID = st.session_state.userID
    query_date_start = col2.date_input('查询开始时间', value=datetime.date.today() - datetime.timedelta(days=1), max_value="today")
    query_date_end = col3.date_input('查询结束时间', value=query_date_start, min_value=query_date_start, max_value="today")
    if st.session_state.userType == 'admin':
        with col1:
            flag_all_user = sac.switch("查询所有用户", value=False, align="start", on_label="On")
    else:
        flag_all_user = False
    with col2:
        flag_approved = sac.switch("仅限已核定工作", value=False, on_label="On")
    if not flag_all_user:
        userID = [query_userID]
        userCName = [query_userCName]
    with col3:
        #dur_time = query_date_end - query_date_start
        chart_type_pack = ['折线图', '中位数图', '旭日图', '矩阵树图', '饼图', '日历热度图']
        if len(userID) == 1:
            chart_type_pack = chart_type_pack + ['柱状图(分组)', '柱状图(堆叠)', '漏斗图']
        chart_type = st.selectbox("图表类型", chart_type_pack, index=1)
    min_value, max_value = 1000, 0
    raws_data, df, fig = [], [], None
    charArea = tab1.empty()
    if chart_type == '折线图':
        # 双Y轴折线图
        fig = go.Figure()
        for index, value in enumerate(userID):
            hot_value, hot_date, temp_value_pack = [], [], []
            sql = f"SELECT task_date, sum(task_score) from clerk_work where task_approved >= {int(flag_approved)} and clerk_id = {value} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' GROUP BY task_date order by task_date"
            result = execute_sql(cur, sql)
            for each in result:
                hot_date.append(each[0])
                hot_value.append(int(each[1]))
                raws_data.append([userCName[index], each[0], int(each[1])])
            temp_value_pack = hot_value
            if temp_value_pack:
                temp_value_pack.sort()
                if temp_value_pack[0] < min_value:
                    min_value = temp_value_pack[0]
                if temp_value_pack[-1] > max_value:
                    max_value = temp_value_pack[-1]
                #st.write(min_value, max_value, hot_value)
                if temp_value_pack[-1] < max_value / 2:
                    yax = 'y2'
                else:
                    yax = 'y'
                fig.add_trace(
                    go.Scatter(name=f"{userCName[index]}",
                                x=hot_date,
                                y=hot_value,
                                mode="markers+lines+text",
                                text=[
                                    format(round(x, 2), ",")
                                    for x in hot_value
                                ],
                                yaxis=yax,
                                textposition="top center"))
        if raws_data:
            fig.update_layout(
                title="工作量",
                xaxis=dict(title="日期"),
                yaxis=dict(title="主轴",
                            rangemode="normal"),
                template="simple_white",
                font=dict(size=st.session_state.chart_font_size),
                yaxis2=dict(
                    title="",
                    overlaying='y',
                    side='right'))
            df = pd.DataFrame(raws_data, columns=["姓名", "日期", "合计分值"])
    elif chart_type.startswith("柱状图"):
        if "分组" in chart_type:
            bar_type = "group"
        elif "堆叠" in chart_type:
            bar_type = "stack"
        else:
            bar_type = "relative"
        for index, value in enumerate(userID):
            sql = f"SELECT task_date, task_group, sum(task_score) from clerk_work where task_approved >= {int(flag_approved)} and clerk_id = {value} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' GROUP BY task_date, task_group order by task_date"
            result = execute_sql(cur, sql)
            for each in result:
                raws_data.append([userCName[index], each[0], each[1], int(each[2])])
        if raws_data:
            df = pd.DataFrame(raws_data, columns=["姓名", "日期", "工作组别", "合计分值"])
            # 使用 Plotly Express 生成分组柱状图
            fig = px.bar(
                df,
                x="日期",
                y="合计分值",
                color="工作组别",
                text="合计分值",
                title="按日期和工作组别统计",
                labels={"合计分值": "总分", "日期": "工作日期", "工作组别": "任务组"},
                barmode=bar_type
            )
            # 调整样式
            fig.update_traces(textposition='outside')
            fig.update_layout(
                xaxis_tickangle=-45,
                template="simple_white",
                font=dict(size=st.session_state.chart_font_size)
            )
    elif chart_type == "日历热度图":
        with tab1:
            with charArea.container(border=True):
                query_date_start = f'{datetime.datetime.now().year}-01-01'
                query_date_end = f'{datetime.datetime.now().year}-12-31'
                cal_data = []
                sql = f"SELECT task_date, sum(task_score) from clerk_work where StationCN = '{st.session_state.StationCN}' and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' group by task_date"
                rows = execute_sql(cur, sql)
                for row in rows:
                    raws_data.append([row[0], int(row[1])])
                    temp = {'value': int(row[1]), 'day': row[0].strftime("%Y-%m-%d")}
                    cal_data.append(temp)
                if raws_data:
                    df = pd.DataFrame(raws_data, columns=["日期", "合计分值"])
                    calendar_chart = {
                        "data": cal_data,
                        "layout": {
                            "title": "日历热度图",
                            "type": "calendar",
                            "height": 500,
                            "from": query_date_start,
                            "to": query_date_end,
                            "emptyColor": "#eeeeee",
                            "colors": ["#61cdbb", "#97e3d5", "#e8c1a0", "#f47560"],
                            "margin": {"top": 40, "right": 40, "bottom": 40, "left": 40},
                            "yearSpacing": 40,
                            "monthBorderColor": "#ffffff",
                            "dayBorderWidth": 2,
                            "dayBorderColor": "#ffffff",
                            "legends": [
                                {
                                    "anchor": "bottom-right",
                                    "direction": "row",
                                    "translateY": 36,
                                    "itemCount": 4,
                                    "itemWidth": 42,
                                    "itemHeight": 36,
                                    "itemsSpacing": 14,
                                    "itemDirection": "right-to-left",
                                }
                            ],
                        },
                    }
                    nc.nivo_chart(data=calendar_chart["data"], layout=calendar_chart["layout"])
                else:
                    st.info("未查询到记录")
    elif chart_type == "漏斗图":
        for index, value in enumerate(userID):
            sql = f"SELECT task_date, task_group, sum(task_score) from clerk_work where task_approved >= {int(flag_approved)} and clerk_id = {value} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' GROUP BY task_date, task_group order by sum(task_score) DESC"
            result = execute_sql(cur, sql)
            for each in result:
                raws_data.append([userCName[index], each[0], each[1], int(each[2])])
        if raws_data:
            df = pd.DataFrame(raws_data, columns=["姓名", "日期", "工作组别", "合计分值"])
            # 按工作组别统计总分并按降序排序
            funnel_data = df.groupby("工作组别")["合计分值"].sum().reset_index()
            funnel_data = funnel_data.sort_values(by="合计分值", ascending=False)
            # 生成漏斗图
            fig = px.funnel(
                funnel_data,
                x="合计分值",
                y="工作组别",
                title="工作量(日期合并)",
                labels={"合计分值": "总分", "工作组别": "任务组别"},
                color_discrete_sequence=px.colors.qualitative.Prism
            )
            # 放大图表整体字体
            fig.update_layout(font=dict(size=st.session_state.chart_font_size))
    elif chart_type == "饼图":
        for index, value in enumerate(userID):
            sql = f"SELECT task_group, sum(task_score) from clerk_work where task_approved >= {int(flag_approved)} and clerk_id = {value} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' GROUP BY task_group order by sum(task_score) DESC"
            result = execute_sql(cur, sql)
            for each in result:
                raws_data.append([userCName[index], each[0], int(each[1])])
        if raws_data:
            df = pd.DataFrame(raws_data, columns=["姓名", "工作组别", "合计分值"])
            if len(userID) > 1:
                pie_data = df.groupby("工作组别")["合计分值"].sum().reset_index()
            else:
                pie_data = df.copy()
            # 计算总和
            total = pie_data['合计分值'].sum()
            # 添加百分比列
            pie_data['百分比'] = (pie_data['合计分值'] / total) * 100
            # 保留所有原始条目，不进行合并
            final_data = pie_data.copy()
            fig = px.pie(
                final_data,
                names="工作组别",
                values="合计分值",
                title="工作量(日期合并)",
                hole=0.2,
                hover_data=["合计分值"],
                labels={"合计分值": "总分", "工作组别": "任务组别"},
                color_discrete_sequence=px.colors.qualitative.Prism
            )
            fig.update_traces(textposition='outside', textinfo='percent+label')
            fig.update_layout(showlegend=False, font=dict(size=12))
    elif chart_type == "旭日图":
        for index, value in enumerate(userID):
            # 查询每个用户的任务分值按工作组别汇总
            sql = f"SELECT task_group, sum(task_score) from clerk_work where task_approved >= {int(flag_approved)} and clerk_id = {value} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' GROUP BY task_group order by sum(task_score) DESC"
            result = execute_sql(cur, sql)
            for each in result:
                raws_data.append([userCName[index], each[0], int(each[1])])
        if raws_data:
            # 构造 DataFrame
            df = pd.DataFrame(raws_data, columns=["姓名", "工作组别", "合计分值"])
            if len(userID) > 1:
                sunburst_data = df.groupby(["工作组别", "姓名"], as_index=False)["合计分值"].sum()
            else:
                sunburst_data = df.copy()
            # 绘制旭日图
            fig = px.sunburst(
                sunburst_data,
                path=['工作组别', '姓名'],
                values='合计分值',
                color='合计分值',
                hover_data=['合计分值'],
                color_continuous_scale='Plasma',
                title="工作量（工作组别 → 用户）",
            )
            fig.update_layout(margin=dict(t=50, l=0, r=0, b=0), font=dict(size=st.session_state.chart_font_size))
    elif chart_type == "矩阵树图":
        for index, value in enumerate(userID):
            # 查询每个用户的任务分值按工作组别汇总
            sql = f"SELECT task_group, sum(task_score) from clerk_work where task_approved >= {int(flag_approved)} and clerk_id = {value} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' GROUP BY task_group order by sum(task_score) DESC"
            result = execute_sql(cur, sql)
            for each in result:
                raws_data.append([userCName[index], each[0], int(each[1])])
        if raws_data:
            df = pd.DataFrame(raws_data, columns=["姓名", "工作组别", "合计分值"])
            if len(userID) > 1:
                treemap_data = df.groupby(["工作组别", "姓名"], as_index=False)["合计分值"].sum()
            else:
                treemap_data = df.copy()
            fig = px.treemap(
                treemap_data,
                path=['工作组别', '姓名'],
                values='合计分值',
                color='合计分值',
                color_continuous_scale='Plasma',
                title="工作量（工作组别 → 用户）",
                hover_data={'合计分值': True}
            )
            fig.update_layout(margin=dict(t=50, l=0, r=0, b=0), font=dict(size=st.session_state.chart_font_size))
    elif chart_type == "中位数图":
        sql = f"SELECT clerk_cname, sum(task_score) from clerk_work where StationCN = '{st.session_state.StationCN}' and task_approved >= {int(flag_approved)} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' GROUP BY clerk_cname order by clerk_cname"
        result = execute_sql(cur, sql)
        for each in result:
            raws_data.append([each[0], int(each[1])])
        if raws_data:
            df = pd.DataFrame(raws_data, columns=["姓名", "合计分值"])
            # 计算中位数
            median_score = np.nanmedian(df["合计分值"])
            # 生成柱状图
            fig = px.bar(df, x="姓名", y="合计分值", text_auto=True,
                        title="工作量分值", labels={"姓名": "员工姓名", "合计分值": "总分值"})
            # 添加中位数水平线
            fig.add_shape(type='line',
                        x0=-0.5, x1=len(df) - 0.5,
                        y0=median_score, y1=median_score,
                        line=dict(color='red', dash='dash'))
            fig.update_layout(font=dict(size=st.session_state.chart_font_size))
            # 将中位数标注移到线上方，并调整字体大小
            fig.add_annotation(x=len(df) - 1, y=median_score + 12,  # 向上偏移
                            text=f'中位数: {median_score:.0f}',
                            showarrow=False,
                            font=dict(color='red', size=st.session_state.chart_font_size + 2),
                            xanchor='right',
                            yanchor='bottom')
    if raws_data:
        if chart_type != "日历热度图":
            with tab1:
                with charArea.container(border=True):
                    st.plotly_chart(fig, use_container_width=True)
        tab2.write(df)
    else:
        st.info("未查询到记录")


def input_public_notice():
    st.subheader("公告发布", divider="red")
    col1, col2 = st.columns(2)
    query_date_start = col1.date_input('公告开始时间', value=datetime.date.today(), min_value="today")
    query_date_end = col2.date_input('公告结束时间', value=query_date_start + datetime.timedelta(days=15), min_value=query_date_start, max_value=query_date_start + datetime.timedelta(days=90))
    confirm_btn_public = st.button('发布')
    display_area = st.empty()
    with display_area.container():
        public_text = st.text_area('请输入公告内容')

    st.subheader("公告修改", divider="green")
    col = st.columns(3)
    notice_pack = []
    now = datetime.datetime.now()
    valid_time = now.strftime("%Y-%m-%d")
    sql = f"SELECT notice, start_time, end_time, ID from notices where StationCN = '{st.session_state.StationCN}' and '{valid_time}' >= start_time and '{valid_time}' <= end_time"
    result = execute_sql(cur, sql)
    if result:
        for row in result:
            notice_pack.append(row[0])
        select_notice = col[0].selectbox("请选择公告", notice_pack, index=0)
        select_notice_index = notice_pack.index(select_notice)
        modify_start_time = col[1].date_input("开始时间", value=result[select_notice_index][1])
        modify_end_time = col[2].date_input("结束时间", value=result[select_notice_index][2])
        st.markdown('##### :red[如果删除, 请在修改内容中键入大写的DELETE]')
        modify_public_text = st.text_area(label='修改公告内容', value=select_notice)
        confirm_btn_modify = st.button('修改')
    else:
        confirm_btn_modify = False

    if confirm_btn_public:
        now = datetime.datetime.now()
        pub_time = now.strftime("%Y-%m-%d %H:%M:%S")
        if public_text:
            sql = f"SELECT ID from notices where StationCN = '{st.session_state.StationCN}' and notice = '{public_text}'"
            if not execute_sql(cur, sql):
                sql = f"INSERT INTO notices (notice, start_time, end_time, publisher, pub_time, StationCN) VALUES ('{public_text}', '{query_date_start}', '{query_date_end}', '{st.session_state.userCName}', '{pub_time}', '{st.session_state.StationCN}')"
                execute_sql_and_commit(conn, cur, sql)
                display_area.empty()
                st.success('公告添加成功')
        else:
            st.warning('请输入公告内容')
    elif confirm_btn_modify:
        if modify_public_text != 'DELETE':
            sql = f"UPDATE notices set notice = '{modify_public_text}', start_time = '{modify_start_time}', end_time = '{modify_end_time}' where ID = {result[select_notice_index][3]} and StationCN = '{st.session_state.StationCN}'"
            execute_sql_and_commit(conn, cur, sql)
            st.success("修改成功")
        else:
            sql = f"DELETE FROM notices where ID = {result[select_notice_index][3]} and StationCN = '{st.session_state.StationCN}'"
            execute_sql_and_commit(conn, cur, sql)
            st.success("删除成功")


def public_notice():
    #st.markdown("### <font face='微软雅黑' color=red><center>站内公告</center></font>", unsafe_allow_html=True)
    st.subheader("站内公告", divider="red")

    vlp_folder = './Images/license_plate/user_vlp'
    now = datetime.datetime.now()
    valid_time = now.strftime("%Y-%m-%d")
    sql = f"SELECT notice from notices where StationCN = '{st.session_state.StationCN}' and '{valid_time}' >= start_time and '{valid_time}' <= end_time"
    result = execute_sql(cur, sql)
    if result:
        for index, row in enumerate(result, start=1):
            st.markdown(f'##### :orange[第{index}条. {row[0]}]')
    else:
        st.info("暂无系统公告")
    if st.session_state.vehicle_restrict:
        # 刷新限行信息
        if not st.session_state.vehicle_restrict_info:
            st.session_state.vehicle_restrict_info = get_vehicle_restrict()
        if st.session_state.vehicle_restrict_info:
            st.subheader("您的车辆限行预警", divider="red")
            vehicle_restrict_info = st.session_state.vehicle_restrict_info
            for each in vehicle_restrict_info:
                st.markdown(f'#### {each[:each.find(" <")]}')
                brand_logo = each[each.find(" <") + 2:each.find("> ")]
                vehicle_num = each[each.rfind('> ') + 2:]
                temp = each.replace(f" <{brand_logo}> ", '')
                model_name = temp[temp.find('<') + 1:temp.find('> ')]
                vlp_model_file = f"{vlp_folder}/{st.session_state.userID}_{model_name}.png"
                vlp_brand_file = f"{vlp_folder}/{brand_logo}_{vehicle_num}.png"
                vlp_file = f"{vlp_folder}/{vehicle_num}.png"
                if os.path.exists(vlp_model_file):
                    st.image(vlp_model_file)
                if os.path.exists(vlp_brand_file):
                    st.image(vlp_brand_file)
                elif os.path.exists(vlp_file):
                    st.image(vlp_file)


@st.fragment
def displayBigTime():
    components.html(open("./MyComponentsScript/Clock-Big.txt", "r", encoding="utf-8").read(), height=140)


def aboutLicense():
    st.subheader("License", divider="green")
    st.markdown(open("./LICENSE", "r", encoding="utf-8").read())


def display_weather_gd(city_code):
    weather_area = st.empty()
    with weather_area.container():
        weather_info = gd_weather_now_cache(city_code)
        if weather_info:
            st.markdown(f"#### {weather_info['city']} - 实时天气")
            #st.markdown(f"<div style='text-align:center; font-family:微软雅黑; color:#008080; font-size:18px;'>地区: {weather_info['city']} 天气: {weather_info['weather_icon']} 温度: {weather_info['temperature']} °C {weather_info['temp_icon']}</div>", unsafe_allow_html=True)
            #st.markdown(f"<div style='text-align:center; font-family:微软雅黑; color:#008080; font-size:18px;'>风向: {weather_info['winddirection']} 风力: {weather_info['windpower']} km/h {weather_info['wind_icon']} 湿度: {weather_info['humidity']}% {weather_info['humidity_icon']}</div>", unsafe_allow_html=True)
            wcol = st.columns(3)
            #wcol[0].metric('地区', weather_info['city'])
            wcol[0].metric(label='天气',value=f"{weather_info['weather']} {weather_info['weather_icon']}")
            wcol[1].metric(label='温度', value=f"{weather_info['temperature']} °C {weather_info['temp_icon']}")
            wcol[2].metric(label='湿度', value=f"{weather_info['humidity']}% {weather_info['humidity_icon']}")
            wcol[0].metric(label='风向', value=f"{weather_info['winddirection']}风")
            wcol[1].metric(label='风力', value=f"{weather_info['windpower']} km/s {weather_info['wind_icon']}")
            wcol[2].metric(label='数据更新时间', value=f"{weather_info['reporttime'][5:]} 数据源: 国家气象中心")
            # 设置度量卡片的样式
            style_metric_cards(border_left_color="#8581d9")


@st.cache_data(ttl="10min")
def gd_weather_now_cache(city_code):
    weather_info = get_city_weather(city_code)

    return weather_info


def display_history_weather():
    #st.markdown("### <font face='微软雅黑' color=green><center>历史天气</center></font>", unsafe_allow_html=True)
    st.subheader("历史天气查询", divider="rainbow")
    city_code = st.session_state.hf_city_code
    city_name = st.session_state.cityname
    sql = f"SELECT MIN(weather_date) from weather_history where city_code = '{city_code}'"
    date_result = execute_sql(cur, sql)
    if date_result:
        query_min_date = date_result[0][0]
    else:
        query_min_date = datetime.datetime.now() - datetime.timedelta(days=10)
        query_min_date = query_min_date.strftime('%Y-%m-%d')
    col = st.columns(3)
    query_date_start = col[0].date_input('查询开始时间', value=datetime.date.today() - datetime.timedelta(days=1), min_value=query_min_date, max_value=datetime.date.today() - datetime.timedelta(days=1))
    query_date_end = col[1].date_input('查询结束时间', value=query_date_start, min_value=query_date_start, max_value="today")
    query_temp_max = col[2].number_input('最高温度(°C)', value=35, min_value=25, max_value=50, step=1)
    query_date_convert = str(query_date_start).replace('-', '')
    sql = f"SELECT sunrise, sunset, moonrise, moonset, moonPhase, tempMax, tempMin, humidity, pressure, moon_icon, temp_icon, humidity_icon, temp_hourly, windspeed_hourly, humidity_hourly, weather_icon_hourly, precip_hourly, windscale_hourly FROM weather_history WHERE city_code = '{city_code}' and weather_date = '{query_date_start}'"
    cur.execute(sql)
    result = cur.fetchone()
    if result:
        weather_info = {
                'sunrise': result[0],
                'sunset': result[1],
                'moonrise': result[2],
                'moonset': result[3],
                'moonPhase': result[4],
                'tempMax': result[5],
                'tempMin': result[6],
                'humidity': result[7],
                'pressure': result[8],
                'moon_icon': result[9],
                'temp_icon': result[10],
                'humidity_icon': result[11],
                'temp_hourly': result[12],
                'windspeed_hourly': result[13],
                'humidity_hourly': result[14],
                'weather_icon_hourly': result[15],
                'precip_hourly': result[16],
                'windscale_hourly': result[17]
            }
    else:
        weather_info = get_city_history_weather(city_code, query_date_convert)
        sql = f"INSERT INTO weather_history (weather_date, city_code, city_name, sunrise, sunset, moonrise, moonset, moonPhase, tempMax, tempMin, humidity, pressure, moon_icon, temp_icon, humidity_icon, temp_hourly, weather_hourly, precip_hourly, windir_hourly, windscale_hourly, windspeed_hourly, humidity_hourly, pressure_hourly, weather_icon_hourly) VALUES ('{query_date_start}', '{city_code}', '{city_name}', '{weather_info['sunrise']}', '{weather_info['sunset']}', '{weather_info['moonrise']}', '{weather_info['moonset']}', '{weather_info['moonPhase']}', '{weather_info['tempMax']}', '{weather_info['tempMin']}', '{weather_info['humidity']}', '{weather_info['pressure']}', '{weather_info['moon_icon']}', '{weather_info['temp_icon']}', '{weather_info['humidity_icon']}', '{weather_info['temp_hourly']}', '{weather_info['weather_hourly']}', '{weather_info['precip_hourly']}', '{weather_info['windir_hourly']}', '{weather_info['windscale_hourly']}', '{weather_info['windspeed_hourly']}', '{weather_info['humidity_hourly']}', '{weather_info['pressure_hourly']}', '{weather_info['weather_icon_hourly']}')"
        execute_sql_and_commit(conn, cur, sql)
    data_col = st.columns(2)
    if weather_info:
        with data_col[0]:
            display_area = st.empty()
            with display_area.container(border=False):
                weather_icon_pack, pre_weather_icon, weather_text = weather_info['weather_icon_hourly'].split('/'), '', ''
                precip_pack = [float(value) for value in weather_info['precip_hourly'].split('/')]
                windscale_pack = [int(value) for value in weather_info['windscale_hourly'].split('/')]
                humidity_pack = [int(value) for value in weather_info['humidity_hourly'].split('/')]
                windspeed_pack = [int(value) for value in weather_info['windspeed_hourly'].split('/')]
                for index, value in enumerate(weather_icon_pack):
                    if value != pre_weather_icon:
                        weather_text = weather_text + str(index) + '点 ' + value + ' '
                        pre_weather_icon = value
                windscale_pack.sort(reverse=True)
                windspeed_pack.sort(reverse=True)
                humidity_pack.sort(reverse=True)
                st.markdown(f"#### :green[{query_date_start} 天气记录]")
                st.markdown(f"##### 地区: {city_name} 温度: {weather_info['tempMax']} - {weather_info['tempMin']} °C {weather_info['temp_icon']}")
                st.markdown(f"##### 天气: {weather_text.strip()}")
                st.markdown(f"##### 降水: {int(sum(precip_pack))} mm 最大风力: {windscale_pack[0]} 级 - {windspeed_pack[0]} km/h")
                st.markdown(f"##### 湿度: {humidity_pack[0]} - {humidity_pack[-1]}% {weather_info['humidity_icon']} 气压: {weather_info['pressure']} hPa")
                st.markdown(f"##### 日升: {weather_info['sunrise']} 日落: {weather_info['sunset']}")
                st.markdown(f"##### 月升: {weather_info['moonrise']} 月落: {weather_info['moonset']} 月相: {weather_info['moonPhase']} {weather_info['moon_icon']}")
        st.subheader('')
        st.markdown(f"##### :blue[{query_date_start} 天气曲线]")
        chart_col = st.columns(2)
        with chart_col[0]:
            plot_wind_speed_curve(weather_info['temp_hourly'].split('/'), weather_info['windspeed_hourly'].split('/'))
        with chart_col[1]:
            plot_data_curve(weather_info['humidity_hourly'].split('/'))
    else:
        st.info("未查询到历史天气记录")
    sql = f"SELECT weather_date, tempMax, tempMin, humidity from weather_history WHERE tempMax >= {query_temp_max} and city_code = '{city_code}' and weather_date >= '{query_date_start}' and weather_date <= '{query_date_end}'"
    results = execute_sql(cur, sql)
    with data_col[1]:
        display_area2 = st.empty()
        with display_area2.container(border=False):
            if results:
                st.markdown(f"#### :red[{query_date_start} - {query_date_end} 高温天气记录]")
                st.markdown(f"##### 共计{len(results)}天 分别是:")
                for result in results:
                    st.markdown(f"##### {result[0]} 温度: {result[1]} - {result[2]} °C 湿度: {result[3]}%")
                st.markdown(":gray[数据源: 国家气象中心]")
            else:
                st.info("未查询到高温天气记录")


def plot_wind_speed_curve(hourly_data1, hourly_data2):
    df1 = pd.DataFrame({
        '小时': range(len(hourly_data1)),
        '温度': hourly_data1
    })
    df2 = pd.DataFrame({
        '小时': range(len(hourly_data2)),
        '风力': hourly_data2
    })

    # 使用Plotly Express创建双Y轴折线图
    fig = make_subplots(specs=[[{'secondary_y': True}]])

    # 添加第一个数据集到图表
    fig.add_trace(
        go.Scatter(x=df1['小时'], y=df1['温度'], name='温度', line=dict(color='red')),
        secondary_y=False,
    )

    # 添加第二个数据集到图表
    fig.add_trace(
        go.Scatter(x=df2['小时'], y=df2['风力'], name='风力', line=dict(color='green', dash='dash')),
        secondary_y=True,
    )

    # 设置图表标题和标签
    fig.update_layout(title_text='温度/风力曲线')
    fig.update_xaxes(title_text='小时')
    fig.update_yaxes(title_text='温度 (°C)', secondary_y=False)
    fig.update_yaxes(title_text='风力 (km/h)', secondary_y=True)

    # 显示图表
    chart = st.empty()
    with chart.container(border=True):
        st.plotly_chart(fig, use_container_width=True)


def plot_data_curve(hourly_data):
    # 创建一个包含小时和温度的数据框
    df = pd.DataFrame({
        '小时': range(len(hourly_data)),
        '湿度': hourly_data
    })

    # 使用Plotly Express创建折线图
    fig = px.line(df, x='小时', y='湿度', title='湿度曲线')
    fig.update_yaxes(title_text='湿度 (%)')
    # 显示图表
    chart = st.empty()
    with chart.container(border=True):
        st.plotly_chart(fig, use_container_width=True)


def display_weather_hf(city_code):
    weather_area = st.empty()
    with weather_area.container():
        weather_info = hf_weather_now_cache(city_code)
        city_name = st.session_state.cityname
        if weather_info:
            get_weather_warning(city_code)
            st.markdown(f'#### {city_name} - 实时天气')
            if weather_info['cloud']:
                cloud = weather_info['cloud']
            else:
                cloud = 'N/A'
            if float(weather_info['precip']) > 0.0:
                precip = '☔'
            else:
                precip = '🌂'
            if st.session_state.weather_icon_type:
                st.markdown(f"<font style='font-size:24px; font-weight:bold;'>天气: {weather_info['weather']}</font> {weather_info['weather_icon']}", unsafe_allow_html=True)
            else:
                weather_icon_html = f"""
                    <html>
                    <head>
                        <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/qweather-icons@1.7.0/font/qweather-icons.css">
                    </head>
                    <body>
                        <i class="qi-{weather_info['weather_icon_id']}" style="font-size: 2em;"></i>
                    </body>
                    </html>
                """
                st.markdown(f"<font style='font-size:24px; font-weight:bold;'>天气: {weather_info['weather']}</font> {weather_icon_html}", unsafe_allow_html=True)
            qweather_icon = qweather_logo()
            wcol = st.columns(4)
            wcol[0].markdown(f"<font style='font-size:24px; font-weight:bold;'>🌡️温度: {weather_info['temp']}°C</font>", unsafe_allow_html=True)
            wcol[1].markdown(f"<font style='font-size:24px; font-weight:bold;'>🧘体感温度: {weather_info['feelslike']}°C {weather_info['feelslike_icon']}</font>", unsafe_allow_html=True)
            wcol[2].markdown(f"<font style='font-size:24px; font-weight:bold;'>降水: {weather_info['precip']} mm {precip}</font>", unsafe_allow_html=True)
            wcol[3].markdown(f"<font style='font-size:24px; font-weight:bold;'>大气压强: {weather_info['pressure']} hPa</font>", unsafe_allow_html=True)
            wcol[0].markdown(f"<font style='font-size:24px; font-weight:bold;'>风向: {weather_info['winddir']}</font>", unsafe_allow_html=True)
            wcol[1].markdown(f"<font style='font-size:24px; font-weight:bold;'>风力: {weather_info['windspeed']} km/h {weather_info['wind_icon']}</font>", unsafe_allow_html=True)
            wcol[2].markdown(f"<font style='font-size:24px; font-weight:bold;'>湿度: {weather_info['humidity']}% {weather_info['humidity_icon']}</font>", unsafe_allow_html=True)
            wcol[3].markdown(f"<font style='font-size:24px; font-weight:bold;'>能见度: {weather_info['vis']} km {weather_info['vis_icon']}</font>", unsafe_allow_html=True)
            st.markdown(f"{qweather_icon}数据更新时间: {weather_info['obstime'][5:-6].replace('T', ' ')} 数据源: NMC/ECMWF", unsafe_allow_html=True)


def display_weather_hf_metric(city_code):
    weather_area = st.empty()
    with weather_area.container():
        weather_info = hf_weather_now_cache(city_code)
        city_name = st.session_state.cityname
        if weather_info:
            qweather_icon = qweather_logo()
            if st.session_state.weather_warning:
                get_weather_warning(city_code)
            if st.session_state.weather_aqi:
                get_weather_aqi(city_code)
            st.markdown(f'#### {city_name} - 实时天气')
            if weather_info['cloud']:
                cloud = weather_info['cloud']
            else:
                cloud = 'N/A'
            if float(weather_info['precip']) > 0.0:
                precip = '☔'
            else:
                precip = '🌂'
            weather_info['pf'] = get_weather_precip_future(city_code)
            wcol = st.columns(4)
            wcol[0].metric(label='天气', value=f"{weather_info['weather']} {weather_info['weather_icon']}")
            #wcol[1].metric(label='🌡️温度', value=f"{weather_info['temp']}°C {weather_info['temp_icon']}")
            wcol[1].metric(label='🌡️温度', value=f"{weather_info['temp']}°C")
            wcol[2].metric(label='🧘体感温度', value=f"{weather_info['feelslike']}°C {weather_info['feelslike_icon']}")
            wcol[3].metric(label='降水量(过去1小时)', value=f"{weather_info['precip']} mm {precip}")
            #wcol[1].metric(label='云量', value=f"{cloud}%")
            #wcol[2].metric(label='大气压强', value=f"{weather_info['pressure']} hPa")
            wcol[0].metric(label='风向', value=weather_info['winddir'])
            wcol[1].metric(label='风力', value=f"{weather_info['windspeed']} km/h {weather_info['wind_icon']}")
            wcol[2].metric(label='湿度', value=f"{weather_info['humidity']}% {weather_info['humidity_icon']}")
            if weather_info['pf']:
                weather_info['pf'] = weather_info['pf'].replace('降雨', '')
                if weather_info['pf'].find('，') != -1:
                    pass
                    #weather_info['pf'] = weather_info['pf'][:weather_info['pf'].find('，')]
                if wcswidth(weather_info['pf']) <= 20:
                    wcol[3].metric(label='降水预报', value=f"{weather_info['pf']}")
                else:
                    # 精确匹配 st.metric 样式但缩小 value 字体
                    markdown_str = pf_markdown_str(weather_info['pf'])
                    wcol[3].markdown(markdown_str, unsafe_allow_html=True)
            else:
                wcol[3].metric(label='能见度', value=f"{weather_info['vis']} km {weather_info['vis_icon']}")
            st.markdown(f"{qweather_icon}数据更新时间: {weather_info['obstime'][5:-6].replace('T', ' ')} 数据源: NMC/ECMWF", unsafe_allow_html=True)
            style_metric_cards(border_left_color="#426edd")


def pf_markdown_str(pf):
    # 精确匹配 st.metric 样式但缩小 value 字体
    pf_length = len(pf)
    # 根据长度确定字体大小
    if pf_length <= 4:
        value_font_size = "24px"
    elif pf_length <= 8:
        value_font_size = "22px"
    elif pf_length <= 12:
        value_font_size = "20px"
    else:
        value_font_size = "18px"
    markdown_str = f"""
        <style>
            .custom-metric-container {{
                display: flex;
                flex-direction: column;
                justify-content: center;
                padding: 16px;
                border-radius: 4px;
                background-color: rgb(255, 255, 255);
                box-shadow: rgba(0, 0, 0, 0.05) 0px 2px 4px;
                text-align: left;
                height: 113px;
                border-left: 8px solid rgb(66, 110, 221);
                margin-bottom: 16px;
                border-top: 1px solid lightgray;    /* 仅顶部边框 */
                border-right: 1px solid lightgray;  /* 仅右侧边框 */
                border-bottom: 1px solid lightgray; /* 仅底部边框 */
            }}
            /* 针对小屏幕设备的响应式调整 */
            @media screen and (max-width: 768px) {{
                .custom-metric-container {{
                    height: 90px;
                }}
            }}
            @media screen and (max-width: 480px) {{
                .custom-metric-container {{
                    height: 80px;
                }}
            }}
            @media screen and (min-width: 1000px) {{
                .custom-metric-container {{
                    height: 113px;
                }}
            }}
            @media screen and (min-width: 1200px) {{
                .custom-metric-container {{
                    height: 127px;
                }}
            }}
            .custom-metric-label {{
                font-size: 14px;
                color: rgba(0, 0, 0, 1.0);
                margin-bottom: 4px;
                font-weight: normal;
                margin-left: 18px; /* 向右移动 */
                margin-bottom: 9px; /* 向上移动 */
            }}
            .custom-metric-value {{
                font-size: {value_font_size}; /* 根据内容长度动态调整字体大小 */
                font-weight: 500;
                color: rgb(49, 51, 63);
                margin-left: 18px; /* 向右移动 */
                margin-bottom: 9px; /* 向上移动 */
            }}
        </style>
        <div class="custom-metric-container">
            <div class="custom-metric-label">降水预测</div>
            <div class="custom-metric-value">{pf}</div>
        </div>
    """

    return markdown_str


@st.cache_data(ttl="10min")
def hf_weather_now_cache(city_code):
    weather_info = get_city_now_weather(city_code)

    return weather_info


@st.cache_data(ttl="10min")
def get_weather_warning(city_code):
    weather_warning = get_city_warning_now(city_code)
    if weather_warning:
        st.markdown(f'#### :red[天气预警]')
        warning_id = 1
        for warning in weather_warning:
            warning["severityColor"] = warning["severityColor"].replace("Yellow", "DarkGoldenRod")
            warning_icon_html = f"""
                <html>
                <head>
                    <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/qweather-icons@1.7.0/font/qweather-icons.css">
                </head>
                <body>
                    <i class="qi-{warning['type']}" style="font-size: 1.4em;"></i>
                </body>
                </html>
            """
            st.markdown(f"<font color={warning['severityColor']} style='font-size:22px; font-weight:bold;'>No.{warning_id} {warning['title']}</font> {warning_icon_html}", unsafe_allow_html=True)
            st.markdown(f"<font style='font-size:18px; font-weight:bold;'>{warning['text']}</font>", unsafe_allow_html=True)
            warning_id += 1
        st.divider()


@st.cache_data(ttl="30min")
def get_weather_aqi(city_code):
    sql = f"SELECT Latitude, Longitude from hf_cn_city where Location_ID = {city_code}"
    lat, lon = execute_sql(cur, sql)[0]
    weather_aqi = get_city_aqi(f'{str(lat)[:-2]}_{str(lon)[:-2]}')
    if weather_aqi:
        st.markdown(f'#### :green[空气质量]')
        st.markdown(f"##### {weather_aqi['health']}")
        pollutants_rev = {'PM 2.5': 'PM 2.5', 'PM 10': 'PM 10', 'NO2': '二氧化氮', 'O3': '臭氧', 'SO2': '二氧化硫', 'CO': '一氧化碳'}
        wcol = st.columns(4)
        col_index = 0
        wcol[0].metric(label='空气质量', value=f"AQI: {weather_aqi['aqi']}")
        wcol[1].metric(label='空气等级', value=f"等级: {weather_aqi['category']}")
        wcol[2].metric(label='质量等级', value=f"{weather_aqi['level']} 级")
        if weather_aqi['primaryPollutant']:
            wcol[3].metric(label='主要污染物', value=f"{weather_aqi['primaryPollutant']} {weather_aqi['primaryPollutant_vu']['value']}  {weather_aqi['primaryPollutant_vu']['unit'].replace('m3', 'm³')}")
        pollutants_col = st.columns(6)
        for each in weather_aqi['sub_pollutants']:
            pollutants_col[col_index % 6].metric(label=pollutants_rev[each], value=f"{weather_aqi['sub_pollutants'][each]['value']} {weather_aqi['sub_pollutants'][each]['unit'].replace('m3', 'm³')}")
            col_index += 1
        style_metric_cards(border_left_color=RGBColor(weather_aqi['color']['red'], weather_aqi['color']['green'], weather_aqi['color']['blue']))
        st.divider()


@st.cache_data(ttl="20min")
def get_weather_precip_future(city_code):
    sql = f"SELECT Latitude, Longitude from hf_cn_city where Location_ID = {city_code}"
    lat, lon = execute_sql(cur, sql)[0]
    weather_pf = get_city_pf_weather(f'{str(lat)[:-2]}_{str(lon)[:-2]}')

    return weather_pf


@st.fragment
def displayVisitCounter_static():
    sql = "SELECT pyLM from verinfo where pyFile = 'visitcounter'"
    visitcount = execute_sql(cur, sql)[0][0]

    # 生成徽章样式的访问计数器
    badge_svg = badge(left_text='访  问  次  数', right_text=f'{visitcount}', right_color='brightgreen')
    # 将badge_svg包裹在具有居中和缩放样式的<div>标签中
    centered_and_scaled_badge = f"<div style='display: flex; justify-content: center; transform: scale(1.3);'>{badge_svg}</div>"
    st.markdown(centered_and_scaled_badge, unsafe_allow_html=True)


@st.fragment
def displayAppInfo_static():
    st.markdown(f"<font face='微软雅黑' color=purple size=16><center>**{APPNAME_CN}**</center></font>", unsafe_allow_html=True)
    verinfo, verLM = getVerInfo()
    st.markdown(f"<font face='微软雅黑' size=5><center>软件版本: {int(verinfo / 10000)}.{int((verinfo % 10000) / 100)}.{int(verinfo / 10)} building {verinfo}</center></font>", unsafe_allow_html=True)
    st.markdown(f"<font face='微软雅黑' size=3><center>更新时间: {time.strftime('%Y-%m-%d %H:%M', time.localtime(verLM))}</center></font>", unsafe_allow_html=True)
    update_type, update_content = get_update_content(f"./CHANGELOG.md")
    st.markdown(f"<font face='微软雅黑' color=blue size=4><center>更新内容: {update_type} - {update_content}</center></font>", unsafe_allow_html=True)


def combine_query():
    #st.markdown("### <font face='微软雅黑' color=tear><center>工作量高级查询</center></font>", unsafe_allow_html=True)
    st.subheader("工作量高级查询", divider="violet")
    adm1 = []
    col = st.columns(4)
    clerk_type = ['默认', '值班', '白班']
    btn_query = col[0].button("查询")
    table_name = col[0].selectbox("请选择表", ['固定列表', '工作量表', '城市代码'])
    if table_name == '固定列表':
        table_name = 'gru_pa'
    elif table_name == '工作量表':
        table_name = 'clerk_work'
    elif table_name == '城市代码':
        table_name = 'hf_cn_city'
        sql = "SELECT DISTINCT(Adm1_Name_ZH) FROM hf_cn_city"
        rows = execute_sql(cur, sql)
        for row in rows:
            adm1.append(row[0])
    clerk_cname_pack, task_group_pack = [], []
    sql = f"SELECT userCName from users where StationCN = '{st.session_state.StationCN}' and clerk_pa <> 0 order by login_counter DESC, userCName"
    result = execute_sql(cur, sql)
    for row in result:
        clerk_cname_pack.append({'value': row[0], 'title': row[0]})
    sql = f"SELECT DISTINCT(task_group) from clerk_work where StationCN = '{st.session_state.StationCN}'"
    result = execute_sql(cur, sql)
    for row in result:
        task_group_pack.append({'value': row[0], 'title': row[0]})
    sql = f"SELECT MIN(task_score), MAX(task_score) from clerk_work where StationCN = '{st.session_state.StationCN}'"
    result = execute_sql(cur, sql)
    min_task_score = result[0][0]
    max_task_score = result[0][1]

    if table_name == 'gru_pa':
        config = {
        'fields': {
            'pa_content': {
                'label': '工作内容',
                'type': 'text',
                #'operators': ['contains', 'not_contains', 'starts_with', 'ends_with'],
            },
            'task_group': {
                'label': '工作组别',
                'type': 'select',
                'fieldSettings': {
                'listValues': task_group_pack,
                },
            },
            'task_score': {
                'label': '单项分值',
                'type': 'number',
                'fieldSettings': {
                'min': min_task_score,
                'max': max_task_score,
                'step': 1
                },
                'preferWidgets': ['slider', 'rangeslider'],
            },
            'multi_score': {
                'label': '多倍数',
                'type': 'boolean',
                'operators': ['equal'],
            },
            'default_task': {
                'label': '默认带入',
                'type': 'select',
                'fieldSettings': {
                'listValues': clerk_type,
                },
            },
            'comm_task': {
                'label': '日常工作',
                'type': 'select',
                'fieldSettings': {
                'listValues': clerk_type,
                },
            },
            'pa_share': {
                'label': '共享分值',
                'type': 'boolean',
                'operators': ['equal'],
            },
            'task_type': {
                'label': '共享独占',
                'type': 'boolean',
                'operators': ['equal'],
            }
        }
        }
    elif table_name == 'clerk_work':
        config = {
        'fields': {
            'clerk_cname': {
                'label': '姓名',
                'type': 'select',
                'fieldSettings': {
                'listValues': clerk_cname_pack,
                },
            },
            'task_date': {
                'label': '日期',
                'type': 'date',
                'operators': ['between']
            },
            'clerk_work': {
                'label': '任务内容',
                'type': 'text',
                #'operators': ['contains', 'not_contains', 'starts_with', 'ends_with'],
            },
            'task_group': {
                'label': '工作组别',
                'type': 'select',
                'fieldSettings': {
                'listValues': task_group_pack,
                },
            },
            'task_score': {
                'label': '单项分值',
                'type': 'number',
                'fieldSettings': {
                'min': min_task_score,
                'max': max_task_score,
                'step': 1
                },
                'preferWidgets': ['slider', 'rangeslider'],
            },
            'task_approved': {
                'label': '核定',
                'type': 'boolean',
                'operators': ['equal'],
            },
            'task_multi': {
                'label': '工作量倍数',
                'type': 'number',
                'fieldSettings': {
                'min': 1,
                'max': 10,
                'step': 1
                },
                'preferWidgets': ['slider', 'rangeslider'],
            }
        }
        }
    elif table_name == 'hf_cn_city':
        config = {
        'fields': {
            'Adm1_Name_ZH': {
                'label': '省市名称',
                'type': 'select',
                'fieldSettings': {
                'listValues': adm1,
                },
            },
            'Location_Name_ZH': {
                'label': '查询地区',
                'type': 'text',
            }
        }
        }

    sql_query = condition_tree(
        config,
        return_type='sql',
        placeholder='暂无查询条件',
    )

    if sql_query and btn_query:
        if table_name == 'gru_pa':
            sql = f"SELECT ID, pa_content, pa_score, task_group, multi_score, default_task, comm_task, pa_share, task_type FROM gru_pa where StationCN = '{st.session_state.StationCN}' and {sql_query}"
            sql = sql.replace('true', '1').replace('false', '0').replace("'默认'", '0').replace("'值班'", '1').replace("'白班'", '2')
            result = execute_sql(cur, sql)
            if result:
                df = pd.DataFrame(result, dtype=str)
                df.columns = ["ID", "工作项", "单项分值", "工作组别", "倍数", "默认带入", "日常工作", "共享分值", "共享独占"]
                for index, value in enumerate(result):
                    df.loc[index, "倍数"] = "多倍" if df["倍数"][index] == '1' else "单倍"
                    df.loc[index, "共享分值"] = "共享" if df["共享分值"][index] == '1' else "不共享"
                    df.loc[index, "共享独占"] = "独占" if df["共享独占"][index] == '1' else "非独占"
                    if df["默认带入"][index] == '0':
                        df.loc[index, "默认带入"] = "默认"
                    elif df["默认带入"][index] == '1':
                        df.loc[index, "默认带入"] = "值班"
                    else:
                        df.loc[index, "默认带入"] = "白班"
                    if df["日常工作"][index] == '0':
                        df.loc[index, "日常工作"] = "默认"
                    elif df["日常工作"][index] == '1':
                        df.loc[index, "日常工作"] = "值班"
                    else:
                        df.loc[index, "日常工作"] = "白班"
                st.dataframe(df)
            else:
                st.info("未查询到记录")
        elif table_name == 'clerk_work':
            sql = f"SELECT ID, task_date, clerk_cname, clerk_work, task_score, task_group, task_approved, task_multi FROM clerk_work where StationCN = '{st.session_state.StationCN}' and {sql_query}"
            sql = sql.replace('true', '1').replace('false', '0')
            result = execute_sql(cur, sql)
            if result:
                df = pd.DataFrame(result, dtype=str)
                df.columns = ["ID", "日期", "员工姓名", "工作项", "单项分值", "工作组别", "核定状态", "工作量倍数"]
                for index, value in enumerate(result):
                    df.loc[index, "核定状态"] = "已核定" if df["核定状态"][index] == '1' else "未核定"
                st.dataframe(df)
            else:
                st.info("未查询到记录")
        elif table_name == 'hf_cn_city':
            sql = f"SELECT Location_Name_ZH, Adm1_Name_ZH, Adm2_Name_ZH, Location_ID, AD_code, Latitude, Longitude FROM hf_cn_city where {sql_query}"
            result = execute_sql(cur, sql)
            if result:
                df = pd.DataFrame(result, dtype=str)
                df.columns = ["地区名称", "省市", "市县", "和风城市代码", "高德城市代码", "经度", "纬度"]
                st.dataframe(df)
            else:
                st.info("未查询到记录")


def update_users_setup(param_name, param_value, action_type):
    if action_type == 'update':
        sql = f"UPDATE users_setup SET param_value = {int(param_value)} where param_name = '{param_name}'"
        execute_sql_and_commit(conn, cur, sql)
    elif action_type == 'insert':
        sql = f"INSERT INTO users_setup (userID, userCName, param_name, param_value) VALUES ({st.session_state.userID}, '{st.session_state.userCName}', '{param_name}', {int(param_value)})"
        execute_sql_and_commit(conn, cur, sql)


def users_setup():
    #st.markdown("### <font face='微软雅黑' color=blue><center>个人设置</center></font>", unsafe_allow_html=True)
    st.subheader("个人设置", divider="green")
    user_setup_name, user_setup_name_intro, user_setup_default = init_user_setup_name()
    col_limit = 3
    col = st.columns(col_limit)
    col_index = 0
    for index, value in enumerate(user_setup_name):
        sql = f"SELECT userID, userCName, param_value from users_setup where userID = {st.session_state.userID} and param_name = '{value}'"
        cur.execute(sql)
        result = cur.fetchone()
        col[col_index % col_limit].markdown(f'##### {user_setup_name_intro[index]}')
        if result:
            with col[col_index % col_limit]:
                sac.switch(label='', value=bool(result[2]), key=f'setup_{value}_{result[0]}', align='start', on_label='On', off_label='Off')
            update_users_setup(value, st.session_state[f'setup_{value}_{result[0]}'], 'update')
        else:
            with col[col_index % col_limit]:
                sac.switch(label='', value=bool(user_setup_default[index]), key=f'setup_{value}_{st.session_state.userID}', align='start', on_label='On', off_label='Off')
            update_users_setup(value, st.session_state[f'setup_{value}_{st.session_state.userID}'], 'insert')
        col_index += 1


def refresh_users_setup():
    user_setup_name, user_setup_name_intro, user_setup_default = init_user_setup_name()
    for index, value in enumerate(user_setup_name):
        sql = f"SELECT userID, userCName, param_value from users_setup where userID = {st.session_state.userID} and param_name = '{value}'"
        result = execute_sql(cur, sql)
        if result:
            st.session_state[value] = bool(result[0][2])
        else:
            st.session_state[value] = bool(user_setup_default[index])
            sql = f"INSERT INTO users_setup (userID, userCName, param_name, param_value) VALUES ({st.session_state.userID}, '{st.session_state.userCName}', '{value}', {user_setup_default[index]})"
            execute_sql_and_commit(conn, cur, sql)


def get_city_code():
    st.session_state.cityname = STATION_CITYNAME[st.session_state.StationCN]
    sql = f"SELECT location_ID, AD_code from hf_cn_city where Location_Name_ZH = '{st.session_state.cityname}'"
    result = execute_sql(cur, sql)
    if result:
        st.session_state.hf_city_code = result[0][0]
        st.session_state.gd_city_code = result[0][1]
    else:
        st.error("城市编码获取失败")


def auto_get_history_weather():
    city_code = st.session_state.hf_city_code
    city_name = st.session_state.cityname
    for i in range(1, 11):
        #st.progress(value=i / 10, text=f'正在获取 {city_name} 的第{i}天历史天气数据...')
        query_date = datetime.datetime.now() - datetime.timedelta(days=i)
        query_date = query_date.strftime('%Y-%m-%d')
        sql = f"SELECT ID FROM weather_history WHERE city_code = '{city_code}' and weather_date = '{query_date}'"
        cur.execute(sql)
        result = cur.fetchone()
        if not result:
            weather_info = get_city_history_weather(city_code, str(query_date).replace('-', ''))
            sql = f"INSERT INTO weather_history (weather_date, city_code, city_name, sunrise, sunset, moonrise, moonset, moonPhase, tempMax, tempMin, humidity, pressure, moon_icon, temp_icon, humidity_icon, temp_hourly, weather_hourly, precip_hourly, windir_hourly, windscale_hourly, windspeed_hourly, humidity_hourly, pressure_hourly, weather_icon_hourly) VALUES ('{query_date}', '{city_code}', '{city_name}', '{weather_info['sunrise']}', '{weather_info['sunset']}', '{weather_info['moonrise']}', '{weather_info['moonset']}', '{weather_info['moonPhase']}', '{weather_info['tempMax']}', '{weather_info['tempMin']}', '{weather_info['humidity']}', '{weather_info['pressure']}', '{weather_info['moon_icon']}', '{weather_info['temp_icon']}', '{weather_info['humidity_icon']}', '{weather_info['temp_hourly']}', '{weather_info['weather_hourly']}', '{weather_info['precip_hourly']}', '{weather_info['windir_hourly']}', '{weather_info['windscale_hourly']}', '{weather_info['windspeed_hourly']}', '{weather_info['humidity_hourly']}', '{weather_info['pressure_hourly']}', '{weather_info['weather_icon_hourly']}')"
            execute_sql_and_commit(conn, cur, sql)
        else:
            break


def cal_date(diff_days):
    if diff_days > 0:
        result_date = datetime.datetime.now() + datetime.timedelta(days=diff_days)
    else:
        result_date = datetime.datetime.now() - datetime.timedelta(days=abs(diff_days))
    result_date = result_date.strftime('%Y-%m-%d')

    return result_date


def modify_db(sub_func):
    #st.markdown("### <font face='微软雅黑' color=red><center>数据库操作</center></font>", unsafe_allow_html=True)
    st.subheader(sub_func, divider="red")
    st.markdown("#### ⚠️请谨慎操作, 记录不可恢复")

    if sub_func == "重置PA-Number":
        reset_table_num()
    elif sub_func == "重置工作组别":
        btn_reset_task_group = st.button(label="确认重置工作组别热度")
        if btn_reset_task_group:
            sql = "TRUNCATE TABLE users_task_group_freq"
            execute_sql_and_commit(conn, cur, sql)
            update_users_group_frequency()
            st.success("工作组别热度重置完成")
    elif sub_func == "更新ID初始值":
        btn_reset_table_id = st.button(label="确认更新ID初始值")
        if btn_reset_table_id:
            sql = """
                SELECT table_name
                FROM information_schema.tables
                WHERE table_schema = %s AND table_type = 'BASE TABLE'
            """
            result = execute_sql(cur, sql, params=('gru-pa',))
            if result:
                for table_name in result:
                    reset_auto_increment(table_name[0])
    elif sub_func == "更新PA-Share":
        btn_reset_pa_share = st.button(label="确认更新PA-Share")
        if btn_reset_pa_share:
            sql = "SELECT DISTINCT(share_date) from pa_share order by share_date"
            date_results = execute_sql(cur, sql)
            sql = "TRUNCATE TABLE pa_share"
            execute_sql_and_commit(conn, cur, sql)
            for date_result in date_results:
                update_pa_share(date_result[0])
            st.success("PA-Share更新完成")
    elif sub_func == "更新固定分值":
        btn_update_fixed_score = st.button(label="更新固定分值", type='primary')
        if btn_update_fixed_score:
            st.button(label="确认更新", type='secondary', on_click=update_fixed_score)
    elif sub_func == "组别名称修改":
        edit_group_name()
    elif sub_func == "分组内容调整":
        modify_task_group()
    elif sub_func == "工作内容修改":
        edit_task_content()
    elif sub_func == "数据库备份":
        btn_backup = st.button(label="开始备份")
        if btn_backup:
            backup_file = f"./MySQL_Backup/GRU-PA-MySQL_Backup_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.sql"
            cmd = f'mysqldump --defaults-file=.mysql.cnf gru-pa > {backup_file}'
            os.system(cmd)
            if os.path.exists(backup_file):
                backup_file_size = round(os.path.getsize(backup_file) / 1024, 1)
                st.success(f"{backup_file[backup_file.rfind('/') + 1:-4]} 文件大小: {backup_file_size}KB 数据库备份完成")
                with open(backup_file, "rb") as file:
                    content = file.read()
                file.close()
                buttonDL = st.download_button("点击下载", content, file_name=backup_file[backup_file.rfind("/") + 1:], icon=":material/download:", type="secondary")
                if buttonDL:
                    st.toast("文件已下载至你的默认目录")
            else:
                st.error("数据库备份失败")


def reset_auto_increment(table_name):
    """
    重置指定表的自增 ID 初始值为当前最大 ID + 1
    :param table_name: 表名
    """
    try:
        # 查询最大 ID
        sql = f"SELECT MAX(id) FROM {table_name}"
        result = execute_sql(cur, sql)
        max_id = result[0][0] if result and result[0][0] is not None else 0

        # 设置自增为最大 ID + 1
        next_auto_inc = max_id + 1
        alter_sql = f"ALTER TABLE {table_name} AUTO_INCREMENT = {next_auto_inc}"
        if execute_sql_and_commit(conn, cur, alter_sql):
            st.success(f"✅ 表 {table_name} 自增ID已重置为 {next_auto_inc}")
        else:
            st.error(f"❌ 表 {table_name} 自增ID重置失败, 请检查表是否存在或权限是否足够")
    except Exception as e:
        st.error(f"⚠️ 表 {table_name} 重置发生异常：{e}")


def update_fixed_score():
    sql = f"SELECT pa_content, pa_score from gru_pa where pa_share = 0 and StationCN = '{st.session_state.StationCN}'"
    rows = execute_sql(cur, sql)
    for row in rows:
        sql = f"UPDATE clerk_work SET task_score = {row[1]} * task_multi where pa_content = '{row[0]}' and StationCN = '{st.session_state.StationCN}'"
        execute_sql_and_commit(conn, cur, sql)
    st.success("固定分值更新成功")


def update_users_group_frequency():
    sql = f"SELECT DISTINCT(task_group) from gru_pa where StationCN = '{st.session_state.StationCN}'"
    rows = execute_sql(cur, sql)
    for row in rows:
        sql = f"SELECT ID from users_task_group_freq where userID = {st.session_state.userID} and task_group = '{row[0]}'"
        if not execute_sql(cur, sql):
            sql = f"INSERT INTO users_task_group_freq (userID, userCName, task_group) VALUES ({st.session_state.userID}, '{st.session_state.userCName}', '{row[0]}')"
            execute_sql_and_commit(conn, cur, sql)
        sql = f"SELECT count(ID) from clerk_work where clerk_ID = {st.session_state.userID} and task_group = '{row[0]}'"
        result = execute_sql(cur, sql)
        sql = f"UPDATE users_task_group_freq SET task_group_freq = {result[0][0]} where userID = {st.session_state.userID} and userCName = '{st.session_state.userCName}' and task_group = '{row[0]}'"
        execute_sql_and_commit(conn, cur, sql)


def get_users_portrait():
    st.subheader("录入人脸数据", divider="green")
    st.markdown(":red[仅限本人使用， 否则识别率会大幅降低]")
    col1, col2 = st.columns(2)
    if st.session_state.userType == 'admin':
        temp_userID, temp_userCName = [], []
        sql = f"SELECT userID, userCName from users where StationCN = '{st.session_state.StationCN}' order by login_counter DESC, userCName"
        rows = execute_sql(cur, sql)
        for row in rows:
            temp_userID.append(row[0])
            temp_userCName.append(row[1])
        img_userCName = col1.selectbox("请选择录入用户", temp_userCName)
        img_userID = temp_userID[temp_userCName.index(img_userCName)]
    else:
        img_userCName = st.session_state.userCName
        img_userID = st.session_state.userID
        st.markdown(f'##### 当前用户: :blue[{img_userCName}]')
    img_file_buffer = st.camera_input("获取人脸图像", width=800)

    if img_file_buffer is not None:
        # To read image file buffer as bytes:
        bytes_data = img_file_buffer.getvalue()
        # Check the type of bytes_data:
        # Should output: <class 'bytes'>
        btn_save_picture = st.button("生成人脸数据")
        pic_file = f"./ID_Photos/{img_userID}_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.jpg"
        if btn_save_picture:
            with open(pic_file, "wb") as f:
                f.write(bytes_data)
            if os.path.exists(pic_file):
                update_face_data(pic_file.replace('./ID_Photos/', './ID_Photos\\'))
                st.toast(f"{img_userCName} 人脸数据获取成功!")
            else:
                st.toast(f"{img_userCName} 人脸数据获取失败!")


@st.fragment
def camera_capture(stationCN):
    st.subheader("人脸识别", divider="green")
    st.markdown('请点击:red[Take Photo]获取人脸图像, 人脸识别后照片会自动销毁, 请放心使用')
    img_file_buffer = st.camera_input("获取人脸图像", width=800)
    if img_file_buffer is not None:
        # To read image file buffer as bytes:
        bytes_data = img_file_buffer.getvalue()
        # Check the type of bytes_data:
        # Should output: <class 'bytes'>
        cap_file = f"./ID_Photos/snapshot_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.jpg"
        with open(cap_file, "wb") as f:
            f.write(bytes_data)
        if os.path.exists(cap_file):
            result = face_login_webrtc(stationCN, cap_file)
            os.remove(cap_file)
            if result:
                login_init(result)
                st.rerun()


def fr_web_rtc():
    st.subheader("人脸识别", divider="green")
    st.markdown('请点击:red[开始]进行人脸识别')
    face_recog_result = None
    webrtc_ctx = webrtc_streamer(
        key="video-sendonly",
        mode=WebRtcMode.SENDONLY,
        media_stream_constraints={"video": True},
        video_html_attrs={
        "style": {"width": "50%", "margin": "0 auto", "border": "0px black solid"},
        "controls": False,
        "autoPlay": True,
        },
        translations={
            "start": "开始",
            "stop": "停止",
            "select_device": "设备选择",
            "media_api_not_available": "无法使用媒体API的环境",
            "device_ask_permission": "请允许访问媒体设备",
            "device_not_available": "无法使用媒体设备",
            "device_access_denied": "访问媒体设备被拒绝",
        },
        # WebRTC 配置中使用STUN服务器
        rtc_configuration={
            "iceServers": [{"urls": ["stun:stun.l.google.com:19302"]}]
        }
    )

    image_place = st.empty()

    while True:
        if webrtc_ctx.video_receiver:
            video_frame = webrtc_ctx.video_receiver.get_frame(timeout=1)
            img_rgb = video_frame.to_ndarray(format="rgb24")
            image_place.image(img_rgb)
            face_recog_result = face_login_webrtc(st.session_state.StationCN, img_rgb)
            if face_recog_result:
                break
        else:
            break

    return face_recog_result


@st.fragment
def face_recognize_verify(stationCN):
    st.subheader("人脸识别验证", divider="rainbow")
    st.markdown('请点击:red[Take Photo]获取人脸图像, 识别后请点击:blue[Clear photo]恢复捕捉')
    col = st.columns(5)
    sql = "SELECT param_value from users_setup where param_name = 'face_tolerance'"
    cur.execute(sql)
    tolerance_now = round(cur.fetchone()[0] / 100, 2)
    tolerance = col[0].number_input("请输入测试容差值 (越低越严格)", min_value=0.2, max_value=0.8, value=0.45, step=0.01, help='输入范围0.2-0.8, 越低越严格')
    col[1].markdown(f"系统当前值: {tolerance_now}")
    flag_update = col[2].checkbox("更新容差值", False)
    img_col = st.columns(2)
    img_file_buffer = img_col[0].camera_input("获取人脸图像", width=800)
    if img_file_buffer is not None:
        st.info(f"容差值为:{round(tolerance, 2)}，人脸识别中...")
        info_col = st.columns(2)
        # To read image file buffer as bytes:
        bytes_data = img_file_buffer.getvalue()
        # Check the type of bytes_data:
        # Should output: <class 'bytes'>
        cap_file = f"./ID_Photos/snapshot_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.jpg"
        with open(cap_file, "wb") as f:
            f.write(bytes_data)
        if os.path.exists(cap_file):
            all_id_distance = face_recognize_webrtc(stationCN, cap_file, tolerance, False)
            if os.path.exists(cap_file):
                os.remove(cap_file)
            if all_id_distance:
                cap_file_point = f'{cap_file[:-4]}_point.jpg'
                if os.path.exists(cap_file_point):
                    with img_col[1]:
                        st.write('面部识别点')
                        st.image(cap_file_point, use_container_width=True)
                col_index = 0
                if flag_update:
                    sql = f"UPDATE users_setup set param_value = {int(round(tolerance, 2) * 100)} where param_name = 'face_tolerance'"
                    execute_sql_and_commit(conn, cur, sql)
                for user_id_distance in all_id_distance:
                    if user_id_distance[0] <= round((100 - st.session_state.min_distance) / 100, 2):
                        sql = f"SELECT userID, userCName, StationCN from users where userID = {user_id_distance[1]}"
                        cur.execute(sql)
                        result = cur.fetchone()
                        info_col[col_index % 2].markdown(f'##### 用户: {result[1]} 站室: {result[2]} 相似度: {round((1 - user_id_distance[0]) * 100, 1)}%')
                        sql = f"SELECT img_filename, upload_time from users_face_data where userID = '{user_id_distance[1]}' and file_hash = '{user_id_distance[2]}' and img_filename is not null"
                        cur.execute(sql)
                        img_file_result = cur.fetchone()
                        if img_file_result:
                            if os.path.exists(img_file_result[0]):
                                info_col[col_index % 2].image(img_file_result[0], caption=f'上传时间:{img_file_result[1]}', use_container_width=True)
                            else:
                                info_col[col_index % 2].write(f"图像文件: {img_file_result[0]} 不存在")
                        else:
                            info_col[col_index % 2].write(f"图像文件不存在")
                        col_index += 1
            else:
                st.markdown('##### 未识别出任何用户!')


@st.fragment
def system_setup():
    #st.markdown("### <font face='微软雅黑' color=blue><center>系统设置</center></font>", unsafe_allow_html=True)
    st.subheader("系统设置", divider="red")
    btn_system_setup_update = st.button("更新系统设置")
    col_limit = 3
    col = st.columns(col_limit)
    col_index = 0
    sql = f"SELECT param_value, param_name, userCName from users_setup where userID = -1 order by ID"
    results = execute_sql(cur, sql)
    for each in results:
        affix_info = ''
        if each[1] == 'max_deduct_score':
            min_value, max_value, step_value = -400, -10, 10
        elif each[1] == 'chart_font_size':
            min_value, max_value, step_value = 10, 20, 1
        elif each[1] == 'md_task_days':
            min_value, max_value, step_value = 28, 31, 1
        elif each[1] == 'max_rev_days':
            min_value, max_value, step_value = 14, 60, 2
        elif each[1] == 'min_distance':
            min_value, max_value, step_value, affix_info = 50, 90, 2, '%'
        st.session_state[each[1]] = col[col_index % col_limit].number_input(label=f'{each[2]}{affix_info}', value=each[0], min_value=min_value, max_value=max_value, step=step_value)
        col_index += 1
    if btn_system_setup_update:
        for each in results:
            if each[0] != st.session_state[each[1]]:
                sql = f"UPDATE users_setup set param_value = {st.session_state[each[1]]} where param_name = '{each[1]}' and userID = -1"
                execute_sql_and_commit(conn, cur, sql)
                st.success(f'{each[2]} 更新成功')


def get_system_setup():
    sql = f"SELECT param_value, param_name from users_setup where userID = -1 order by ID"
    results = execute_sql(cur, sql)
    for each in results:
        st.session_state[each[1]] = each[0]


def init_user_setup_name():
    key_name_pack, key_intro_pack, default_value_pack = [], [], []
    sql = "SELECT key_name, key_intro, default_value from users_setup_template order by key_name, ID"
    results = execute_sql(cur, sql)
    for each in results:
        key_name_pack.append(each[0])
        key_intro_pack.append(each[1])
        default_value_pack.append(each[2])

    return key_name_pack, key_intro_pack, default_value_pack


def init_task_group_icon():
    group_icon = {}
    sql = "SELECT key_name, key_value from icons where icon_type = 'task_group' order by ID"
    results = execute_sql(cur, sql)
    for each in results:
        group_icon[each[0]] = each[1]

    return group_icon


@st.fragment
def modify_task_group():
    org_task_group, item_pack = [], []
    sql = f"SELECT DISTINCT(task_group) from gru_pa where StationCN = '{st.session_state.StationCN}'"
    rows = execute_sql(cur, sql)
    for row in rows:
        org_task_group.append(row[0])
    col1, col2 = st.columns(2)
    org_selected = col1.selectbox('原工作组别', org_task_group, index=0)
    target_selected = col2.selectbox('目标工作组别', org_task_group, index=0)
    btn_modify = col1.button('调整组别')
    if org_selected != target_selected:
        sql = f"SELECT pa_content, ID from gru_pa where task_group = '{org_selected}' and StationCN = '{st.session_state.StationCN}'"
        rows = execute_sql(cur, sql)
        for row in rows:
            item_pack.append(f'{row[0]} ID:{row[1]}')
        moved_pack = sac.transfer(items=item_pack, label='工作组别调整', titles=['项目'], reload=True, align='center', search=True, pagination=True, use_container_width=True)
        if btn_modify and moved_pack:
            for each in moved_pack:
                modify_id = each[each.rfind('ID:') + 3:].strip()
                modify_content = each[:each.rfind('ID:')]
                sql = f"UPDATE gru_pa SET task_group = '{target_selected}' where ID = {modify_id}"
                execute_sql_and_commit(conn, cur, sql)
                sql = f"UPDATE clerk_work SET task_group = '{target_selected}' where clerk_work = '{modify_content}' and StationCN = '{st.session_state.StationCN}'"
                execute_sql_and_commit(conn, cur, sql)
                st.success(f":violet[{org_selected}] 中的 :orange[{modify_content}] 已经调整至 :blue[{target_selected}]")


@st.fragment
def edit_group_name():
    org_task_group = []
    sql = f"SELECT DISTINCT(task_group) from gru_pa where StationCN = '{st.session_state.StationCN}'"
    rows = execute_sql(cur, sql)
    for row in rows:
        org_task_group.append(row[0])
    col1, col2 = st.columns(2)
    group_selected = col1.selectbox('工作组别', org_task_group, index=0)
    change_txt = col2.text_input('修改为', group_selected)
    if group_selected != change_txt:
        btn_change = col1.button('确认修改')
        if btn_change:
            sql = f"UPDATE gru_pa set task_group = '{change_txt}' where task_group = '{group_selected}' and StationCN = '{st.session_state.StationCN}'"
            execute_sql_and_commit(conn, cur, sql)
            sql = f"UPDATE clerk_work set task_group = '{change_txt}' where task_group = '{group_selected}' and StationCN = '{st.session_state.StationCN}'"
            execute_sql_and_commit(conn, cur, sql)
            sql = f"UPDATE icons set key_name = '{change_txt}' where key_name = '{group_selected}' and icon_type = 'task_group'"
            execute_sql_and_commit(conn, cur, sql)
            sql = "TRUNCATE TABLE users_task_group_freq"
            execute_sql_and_commit(conn, cur, sql)
            update_users_group_frequency()
            st.success(f'[{group_selected}] 已修改为 [{change_txt}]')


@st.fragment
def edit_task_content():
    org_task_group, org_task_content = [], []
    sql = f"SELECT DISTINCT(task_group) from gru_pa where StationCN = '{st.session_state.StationCN}'"
    rows = execute_sql(cur, sql)
    for row in rows:
        org_task_group.append(row[0])
    group_selected = st.selectbox('工作组别', org_task_group, index=0)
    sql = f"SELECT pa_content from gru_pa where task_group = '{group_selected}' and StationCN = '{st.session_state.StationCN}'"
    rows = execute_sql(cur, sql)
    for row in rows:
        org_task_content.append(row[0])
    change_selected = st.selectbox('工作内容', org_task_content, index=0)
    st.markdown(f"#### :red[修改内容为DELETE表示删除该项工作]")
    change_txt = st.text_input('修改为', change_selected)
    if change_selected != change_txt:
        btn_change = st.button('确认修改')
        if btn_change:
            if change_txt.upper() != 'DELETE':
                sql = f"UPDATE gru_pa set pa_content = '{change_txt}' where task_group = '{group_selected}' and pa_content = '{change_selected}' and StationCN = '{st.session_state.StationCN}'"
                execute_sql_and_commit(conn, cur, sql)
                sql = f"UPDATE clerk_work set clerk_work = '{change_txt}' where task_group = '{group_selected}' and clerk_work = '{change_selected}' and StationCN = '{st.session_state.StationCN}'"
                execute_sql_and_commit(conn, cur, sql)
                sql = f"UPDATE pa_share set pa_content = '{change_txt}' where pa_content = '{change_selected}' and StationCN = '{st.session_state.StationCN}'"
                execute_sql_and_commit(conn, cur, sql)
                st.success(f'[{change_selected}] 已修改为 [{change_txt}]')
            elif change_txt == 'DELETE':
                sql = f"DELETE from gru_pa where pa_content = '{change_selected}' and task_group = '{group_selected}' and StationCN = '{st.session_state.StationCN}'"
                execute_sql_and_commit(conn, cur, sql)
                sql = f"DELETE from clerk_work where clerk_work = '{change_selected}' and task_group = '{group_selected}' and StationCN = '{st.session_state.StationCN}'"
                execute_sql_and_commit(conn, cur, sql)
                sql = f"DELETE from pa_share where pa_content = '{change_selected}' and StationCN = '{st.session_state.StationCN}'"
                execute_sql_and_commit(conn, cur, sql)
                st.success(f'[{change_selected}] 已删除!')
            elif change_txt == 'delete':
                st.info(f'如果想删除请键入大写 :red[DELETE]')

def qweather_logo():
    qweather_icon = f"""
        <html>
        <head>
            <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/qweather-icons@1.7.0/font/qweather-icons.css">
        </head>
        <body>
            <i class="qi-qweather" style="font-size: 1em;"></i>
        </body>
        </html>
    """

    return qweather_icon


def get_vehicle_restrict():
    restrict_info, gen_vp_pack, brand_logo_pack, vehicle_model_pack, userID_pack = [], [], [], [], []
    for i in range(2):
        query_date = cal_date(i)
        if is_workday(datetime.date.fromisoformat(query_date)):
            query_wor = datetime.date.fromisoformat(query_date).weekday() + 1
            sql = f"SELECT license_plate, userCName, vehicle_brand, vehicle_model from vehicle_info where userID = {st.session_state.userID} and StationCN = '{st.session_state.StationCN}' and vehicle_type = 0"
            results = execute_sql(cur, sql)
            if results:
                for result in results:
                    if result[0][-1].lower() == 'x':
                        user_last_plate = 0
                    else:
                        user_last_plate = result[0][-1]
                    sql = f"SELECT ID from vehicle_restrict where wor = {query_wor} and tail_num = {user_last_plate} and start_time <= '{query_date}' and end_time >= '{query_date}' and StationCN = '{st.session_state.StationCN}'"
                    if execute_sql(cur, sql):
                        if i > 0:
                            restrict_info.append(f':orange[明日限行] <{result[2]}> <{result[3]}> {result[0]}')
                        else:
                            restrict_info.append(f'今日限行 <{result[2]}> <{result[3]}> {result[0]}')
    sql = f"SELECT license_plate, userCName, vehicle_brand, vehicle_model, userID from vehicle_info where StationCN = '{st.session_state.StationCN}' and vehicle_type = 0"
    results = execute_sql(cur, sql)
    if results:
        for result in results:
            if not os.path.exists(f"./Images/license_plate/user_vlp/{result[0]}.png"):
                gen_vp_pack.append(result[0])
                brand_logo_pack.append(result[2])
                vehicle_model_pack.append(result[3])
                userID_pack.append(result[4])
    if gen_vp_pack:
        create_plate_image(gen_vp_pack, brand_logo_pack, vehicle_model_pack, userID_pack)
    if restrict_info:
        return restrict_info

    return None


def update_vehicle_restrict(d1, d2, diff_days):
    sql = f"SELECT tail_num, wor, StationCN from vehicle_restrict where ID < 11 and StationCN = '{st.session_state.StationCN}' order by ID"
    rows = execute_sql(cur, sql)
    for row in rows:
        wor = row[1] + diff_days
        if wor > 5:
            wor = wor - 5
        sql = f"INSERT INTO vehicle_restrict (tail_num, start_time, end_time, wor, StationCN) VALUES ({row[0]}, '{d1}', '{d2}', {wor}, '{row[2]}')"
        #execute_sql_and_commit(conn, cur, sql)


def bonus_scene():
    st.subheader(":orange[站内车库]", divider="rainbow")
    vlp_folder = './Images/license_plate/user_vlp'
    col = st.columns(4)
    col_index = 0
    sql = f"SELECT userID, license_plate, vehicle_brand, vehicle_model FROM vehicle_info where StationCN = '{st.session_state.StationCN}' order by userID"
    rows = execute_sql(cur, sql)
    for row in rows:
        vlp_vehicle_file = f"{vlp_folder}/{row[0]}_{row[1]}_{row[2]}_{row[3]}.png"
        vlp_file = f"{vlp_folder}/{row[2]}_{row[1]}.png"
        if os.path.exists(vlp_vehicle_file):
            col[col_index % 4].image(vlp_vehicle_file)
        if os.path.exists(vlp_file):
            col[col_index % 4].image(vlp_file)
        if os.path.exists(vlp_file) or os.path.exists(vlp_vehicle_file):
            col[col_index % 4].divider()
            col_index += 1


def temp_func():
    pass


global APPNAME_CN, APPNAME_EN, WEATHERICON, STATION_CITYNAME
APPNAME_CN = "站室绩效考核系统GRU-PA"
APPNAME_EN = "GRU-PA"
STATION_CITYNAME = {'北京站': '顺义', '天津站': '滨海新区', '总控室': '滨海新区', '调控中心': '滨海新区', '武清站': '武清'}
conn = get_connection()
cur = conn.cursor()
st.logo(image="./Images/logos/GRU-PA-logo.png", icon_image="./Images/logos/GRU-PA-logo.png", size="large")
selected = None

if "logged_in" not in st.session_state:
    if 'Windows' in st.context.headers['User-Agent']:
        st.session_state.is_moble = False
    else:
        st.session_state.is_moble = True
    if st.session_state.is_moble:
        with open("./MyComponentsScript/mobile-style.css", "r", encoding="utf-8") as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
    update_face_data()
    clean_snapshot()
    st.session_state.client_local = True if st.context.headers['host'].startswith('localhost') else False
    st.session_state.logged_in = False
    st.session_state.login_webrtc = False
    login()
elif st.session_state.login_webrtc:
    result = fr_web_rtc()
    if result:
        login_init(result)
        st.session_state.login_webrtc = False
        st.rerun()
elif st.session_state.logged_in:
    with st.sidebar:
        st.markdown(f"<font face='微软雅黑' color=green size=4><center>**当前用户: {st.session_state.userCName}**</center></font>", unsafe_allow_html=True)
        #st.markdown(f'### :green[当前用户:] :orange[{st.session_state.userCName}]')
        #displaySmallTime()
        #displaySmallClock()
        if st.session_state.menu_index == 0:
            notice_icon = 'megaphone-fill'
        else:
            notice_icon = 'megaphone'
        if st.session_state.userType == "admin":
            if st.session_state.userID == 1 and st.session_state.StationCN == "北京站":
                system_da = True
            else:
                system_da = False
            selected = sac.menu([
                sac.MenuItem('公告', icon=notice_icon),
                sac.MenuItem('主页', icon='house'),
                sac.MenuItem('功能', icon='columns-gap', children=[
                    sac.MenuItem('工作量批量录入', icon='list-task'),
                    sac.MenuItem('工作量手工录入', icon='journal-plus'),
                    sac.MenuItem('工作减分项录入', icon='journal-minus'),
                    sac.MenuItem('记录修改', icon='journal-medical'),
                    sac.MenuItem('统计查询及导出', icon='clipboard-data'),
                    sac.MenuItem('趋势图', icon='bar-chart-line'),
                    sac.MenuItem('数据检查与核定', icon='check2-all'),
                    sac.MenuItem('高级查询', icon='search'),
                    sac.MenuItem('历史天气查询', icon='cloud-sun'),
                    sac.MenuItem('公告发布', icon='journal-arrow-up'),
                ]),
                sac.MenuItem('数据库操作', icon='database-check', children=[
                    sac.MenuItem('重置PA-Number', icon='bootstrap-reboot'),
                    sac.MenuItem('重置工作组别', icon='vignette'),
                    sac.MenuItem('更新ID初始值', icon='database-exclamation'),
                    sac.MenuItem('更新PA-Share', icon='list-check'),
                    sac.MenuItem('更新固定分值', icon='database-up'),
                    sac.MenuItem('工作组别调整', icon='collection', children=[
                        sac.MenuItem('组别名称修改', icon='pencil-square'),
                        sac.MenuItem('分组内容调整', icon='text-indent-left'),
                        sac.MenuItem('工作内容修改', icon='card-text'),
                    ]),
                    sac.MenuItem('数据库备份', icon='backpack4'),
                ], disabled=not system_da),
                sac.MenuItem('设置', icon='gear', children=[
                    sac.MenuItem('个人设置', icon='sliders2'),
                    sac.MenuItem('系统设置', icon='sliders2-vertical', disabled=not system_da),
                    sac.MenuItem('录入人脸数据', icon='person-bounding-box'),
                    sac.MenuItem('人脸识别测试', icon='person-video3'),
                ]),
                sac.MenuItem('账户', icon='person-gear', children=[
                    sac.MenuItem('密码修改', icon='key'),
                    sac.MenuItem('密码重置', icon='person-gear'),
                    sac.MenuItem('登出', icon='box-arrow-right'),
                ]),
                sac.MenuItem('关于', icon='info-circle', children=[
                    sac.MenuItem('更新日志', icon='h-square'),
                    sac.MenuItem('自述', icon='clipboard'),
                    sac.MenuItem('彩蛋', icon='bootstrap'),
                    sac.MenuItem('关于', icon='book'),
                    sac.MenuItem('许可证', icon='card-text'),
                ]),
            ], open_index=[1, 2], index=st.session_state.menu_index)
        elif st.session_state.userType == "user":
            selected = sac.menu([
                sac.MenuItem('公告', icon=notice_icon),
                sac.MenuItem('主页', icon='house'),
                sac.MenuItem('功能', icon='columns-gap', children=[
                    sac.MenuItem('工作量批量录入', icon='list-task'),
                    sac.MenuItem('工作量手工录入', icon='journal-plus'),
                    sac.MenuItem('记录修改', icon='journal-medical'),
                    sac.MenuItem('统计查询及导出', icon='clipboard-data'),
                    sac.MenuItem('趋势图', icon='bar-chart-line'),
                    sac.MenuItem('历史天气查询', icon='cloud-sun'),
                ]),
                sac.MenuItem('设置', icon='gear', children=[
                    sac.MenuItem('个人设置', icon='sliders2'),
                    sac.MenuItem('录入人脸数据', icon='person-bounding-box'),
                ]),
                sac.MenuItem('账户', icon='person-gear', children=[
                    sac.MenuItem('密码修改', icon='key'),
                    sac.MenuItem('登出', icon='box-arrow-right'),
                ]),
                sac.MenuItem('关于', icon='info-circle', children=[
                    sac.MenuItem('更新日志', icon='h-square'),
                    sac.MenuItem('自述', icon='clipboard'),
                    sac.MenuItem('彩蛋', icon='bootstrap'),
                    sac.MenuItem('关于', icon='book'),
                    sac.MenuItem('许可证', icon='card-text'),
                ]),
            ], open_index=[1, 2], index=st.session_state.menu_index)
        st.divider()
        st.image(f'./Images/badges/{APPNAME_EN}-badge.svg')
        st.image(f'./Images/badges/{APPNAME_EN}-lm-badge.svg')
    if selected == "公告":
        # 更新限行条件
        #update_vehicle_restrict('2025-09-29', '2025-12-28', 1)
        public_notice()
    elif selected == "主页":
        # 刷新个人设置
        refresh_users_setup()
        # 更新版本信息
        updatePyFileinfo()
        verinfo, verLM = getVerInfo()
        app_version = f'{int(verinfo / 10000)}.{int((verinfo % 10000) / 100)}.{verinfo}'
        app_lm = time.strftime('%Y-%m-%d %H:%M', time.localtime(verLM))
        gen_badge(conn, cur, [], 'MySQL', APPNAME_EN, app_version, app_lm)
        if st.session_state.static_show:
            displayAppInfo_static()
        else:
            displayBigTime()
            displayAppInfo(300)
        if st.session_state.weather_show:
            if st.session_state.weather_provider:
                if st.session_state.weather_metric:
                    display_weather_hf_metric(st.session_state.hf_city_code)
                    # 手动测试
                    #display_weather_hf_metric('101050304')
                else:
                    display_weather_hf(st.session_state.hf_city_code)
                    st.header(' ')
            else:
                display_weather_gd(st.session_state.gd_city_code)
        else:
            st.header(' ')
            st.header(' ')
        sac.divider(label='', icon=sac.BsIcon(name='boxes', size=20), align='center')
        if st.session_state.static_show:
            displayVisitCounter_static()
        else:
            displayVisitCounter()
    elif selected == "工作量批量录入":
        task_input()
    elif selected == "工作量手工录入":
        manual_input()
    elif selected == "工作减分项录入":
        deduction_input()
    elif selected == "记录修改":
        task_modify()
    elif selected == "统计查询及导出":
        query_task()
    elif selected == "趋势图":
        gen_chart()
    elif selected == "数据检查与核定":
        check_data()
    elif selected == "高级查询":
        combine_query()
    elif selected == "公告发布":
        input_public_notice()
    elif selected == "历史天气查询":
        display_history_weather()
    elif selected == "重置PA-Number" or selected == "重置工作组别" or selected == "更新ID初始值" or selected == "更新PA-Share" or selected == "更新固定分值" or selected == "组别名称修改" or selected == "分组内容调整" or selected == "工作内容修改" or selected == "数据库备份":
        modify_db(selected)
    elif selected == "个人设置":
        users_setup()
    elif selected == "系统设置":
        system_setup()
    elif selected == "录入人脸数据":
        get_users_portrait()
    elif selected == "人脸识别测试":
        face_recognize_verify(st.session_state.StationCN)
    elif selected == "密码修改":
        changePassword()
    elif selected == "密码重置":
        resetPassword()
    elif selected == "登出":
        logout()
    elif selected == "更新日志":
        changelog()
    elif selected == "自述":
        aboutReadme()
    elif selected == "关于":
        aboutInfo()
    elif selected == "彩蛋":
        bonus_scene()
    elif selected == "许可证":
        aboutLicense()
