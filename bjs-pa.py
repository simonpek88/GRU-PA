# coding utf-8
import datetime
import os
import time

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
import streamlit.components.v1 as components
import streamlit_antd_components as sac
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from streamlit_extras.badges import badge

from commFunc import (execute_sql, execute_sql_and_commit, get_update_content,
                      getUserEDKeys, getVerInfo, updatePyFileinfo)
from mysql_pool import get_connection

# cSpell:ignoreRegExp /[^\s]{16,}/
# cSpell:ignoreRegExp /\b[A-Z]{3,15}\b/g
# cSpell:ignoreRegExp /\b[A-Z]\b/g

@st.fragment
def login():
    # 显示应用名称
    st.markdown(f"<font face='微软雅黑' color=purple size=20><center>**{APPNAME}**</center></font>", unsafe_allow_html=True)

    # 登录表单容器
    login = st.empty()
    with login.container(border=True):
        # 用户编码输入框
        userID = st.text_input("请输入用户编码", placeholder="请输入纯数字用户编码", max_chars=8)
        # 初始化用户姓名
        st.session_state.userCName = ""

        # 用户密码输入框
        userPassword = st.text_input("请输入密码", max_chars=8, placeholder="用户初始密码为1234", type="password", autocomplete="off")

        # 登录按钮
        buttonLogin = st.button("登录")

    # 如果点击了登录按钮
    if buttonLogin:
        # 如果用户编码和密码不为空
        if userID != "" and userPassword != "":
            # 验证用户密码
            verifyUPW = verifyUserPW(userID, userPassword)
            # 如果密码验证成功
            if verifyUPW[0]:
                userPassword = verifyUPW[1]
            sql = f"SELECT userID, userCName, userType, StationCN, clerk_type from users where userID = {userID} and userPassword = '{userPassword}'"
            result = execute_sql(cur, sql)
            if result:
                st.toast(f"用户: {result[0][0]} 姓名: {result[0][1]} 登录成功, 欢迎回来")
                login.empty()
                st.session_state.logged_in = True
                st.session_state.userID = result[0][0]
                st.session_state.userCName = result[0][1]
                st.session_state.userType = result[0][2]
                st.session_state.StationCN = result[0][3]
                st.session_state.clerkType = result[0][4]
                st.session_state.userPwRechecked = False
                sql = "UPDATE verinfo set pyLM = pyLM + 1 where pyFile = 'visitcounter'"
                execute_sql_and_commit(conn, cur, sql)
                updatePyFileinfo()
                st.rerun()
            elif not verifyUPW[0]:
                st.error("登录失败, 请检查用户名和密码, 若忘记密码请联系管理员重置")
        else:
            st.warning("请输入用户编码和密码")


def logout():
    # 关闭游标
    cur.close()
    # 关闭数据库连接
    conn.close()

    # 清除会话状态中的所有键值对
    for key in st.session_state.keys():
        del st.session_state[key]

    # 重新运行当前脚本
    st.rerun()


def verifyUserPW(vUserName, vUserPW):
    st.session_state.userPwRechecked = False
    vUserEncPW = ""
    sql = f"SELECT userPassword from users where userID = {vUserName}"
    pwTable = execute_sql(cur, sql)
    if pwTable:
        vUserEncPW = pwTable[0][0]
        vUserDecPW = getUserEDKeys(vUserEncPW, "dec")
        if vUserPW == vUserDecPW:
            st.session_state.userPwRechecked = True

    return st.session_state.userPwRechecked, vUserEncPW

@st.fragment
def displayBigTimeCircle():
    components.html(open("./MyComponentsScript/Clock-Big-Circle.txt", "r", encoding="utf-8").read(), height=260)


@st.fragment
def displayVisitCounter():
    sql = "SELECT pyLM from verinfo where pyFile = 'visitcounter'"
    visitcount = execute_sql(cur, sql)[0][0]
    countScript = (open("./MyComponentsScript/FlipNumber.txt", "r", encoding="utf-8").read()).replace("visitcount", str(visitcount))
    components.html(countScript, height=40)


@st.fragment
def displaySmallTime():
    components.html(open("./MyComponentsScript/Clock-Small.txt", "r", encoding="utf-8").read(), height=34)


@st.fragment
def displaySmallClock():
    components.html(open("./MyComponentsScript/Clock-Number.txt", "r", encoding="utf-8").read(), height=30)


@st.fragment
def displayAppInfo():
    infoStr = open("./MyComponentsScript/glowintext.txt", "r", encoding="utf-8").read()
    infoStr = infoStr.replace("软件名称", APPNAME)
    verinfo, verLM = getVerInfo()
    infoStr = infoStr.replace("软件版本", f"软件版本: {int(verinfo / 10000)}.{int((verinfo % 10000) / 100)}.{int(verinfo / 10)} building {verinfo}")
    infoStr = infoStr.replace("更新时间", f"更新时间: {time.strftime('%Y-%m-%d %H:%M', time.localtime(verLM))}")
    update_type, update_content = get_update_content(f"./CHANGELOG.md")
    infoStr = infoStr.replace("更新内容", f"更新内容: {update_type} - {update_content}")
    components.html(infoStr, height=340)


def changePassword():
    # 显示密码修改页面标题
    st.write("### :red[密码修改]")
    # 创建一个带有边框的容器
    changePW = st.empty()
    with changePW.container(border=True):
        # 输入原密码
        oldPassword = st.text_input("请输入原密码", max_chars=8, type="password", autocomplete="off")
        # 输入新密码
        newPassword = st.text_input("请输入新密码", max_chars=8, type="password", autocomplete="off")
        # 再次输入新密码以确认
        confirmPassword = st.text_input("请再次输入新密码", max_chars=8, placeholder="请与上一步输入的密码一致", type="password", autocomplete="new-password")
        # 确认修改按钮
        buttonSubmit = st.button("确认修改")

    # 检查原密码是否为空
    if oldPassword:
        # 验证用户原密码
        verifyUPW = verifyUserPW(st.session_state.userID, oldPassword)
        if verifyUPW[0]:
            oldPassword = verifyUPW[1]
        # 构造SQL查询语句，验证用户名和密码是否匹配
        sql = f"SELECT ID from users where userID = {st.session_state.userID} and userPassword = '{oldPassword}'"
        if execute_sql(cur, sql):
            # 检查新密码和确认密码是否填写且一致
            if newPassword and confirmPassword and newPassword != "":
                if newPassword == confirmPassword:
                    # 确认修改按钮是否被点击
                    if buttonSubmit:
                        # 加密新密码
                        newPassword = getUserEDKeys(newPassword, "enc")
                        # 构造SQL更新语句，更新用户密码
                        sql = f"UPDATE users set userPassword = '{newPassword}' where userID = {st.session_state.userID}"
                        # 执行SQL语句并提交
                        execute_sql_and_commit(conn, cur, sql)
                        # 显示密码修改成功提示，并要求重新登录
                        st.toast("密码修改成功, 请重新登录")
                        # 登出用户
                        logout()
                else:
                    # 显示密码不一致的错误信息
                    st.error("两次输入的密码不一致")
            else:
                # 显示新密码未填写的警告信息
                st.warning("请检查新密码")
        else:
            # 显示原密码错误的错误信息
            st.error("原密码不正确")
    else:
        st.warning("原密码不能为空")


@st.fragment
def changelog():
    changelogInfo = open("./CHANGELOG.md", "r", encoding="utf-8").read()
    st.markdown(changelogInfo)


def aboutReadme():
    st.markdown(open("./README.md", "r", encoding="utf-8").read())


def aboutInfo():
    st.subheader("关于本软件", divider="rainbow")
    st.subheader(":blue[Powered by Python and Streamlit]")
    logo1, logo2, logo3, logo4, logo5, logo6 = st.columns(6)
    with logo1:
        st.caption("Python")
        st.image("./Images/logos/python.png")
    with logo2:
        st.caption("Streamlit")
        st.image("./Images/logos/streamlit.png")
    with logo3:
        st.caption("MySQL")
        st.image("./Images/logos/mysql.png")
    with logo4:
        st.caption("Ant Comp")
        st.image("./Images/logos/antd.png")
    with logo5:
        st.caption("Pandas")
        st.image("./Images/logos/pandas.png")
    with logo6:
        st.caption("Plotly")
        st.image("./Images/logos/plotly.png")
    display_pypi()
    st.write("###### :violet[为了获得更好的使用体验, 请使用浅色主题]")
    verinfo, verLM = getVerInfo()
    st.caption(f"Version: {int(verinfo / 10000)}.{int((verinfo % 10000) / 100)}.{int(verinfo / 10)} building {verinfo} Last Modified: {time.strftime('%Y-%m-%d %H:%M', time.localtime(verLM))}")
    sac.divider(align="center", color="blue")


def display_pypi():
    pypi1, pypi2, pypi3, pypi4, pypi5, pypi6 = st.columns(6)
    with pypi1:
        badge(type="pypi", name="streamlit")
    with pypi2:
        badge(type="pypi", name="streamlit_antd_components")
    with pypi3:
        badge(type="pypi", name="pandas")
    with pypi4:
        badge(type="pypi", name="plotly")


def actionResetUserPW(rUserID, rOption1, rOption2, rUserType):
    rInfo = ""

    # 如果 rOption1 为真
    if rOption1:
        # 获取用户加密密钥
        resetPW = getUserEDKeys("1234", "enc")
        # 构建 SQL 更新语句
        sql = f"UPDATE users SET userPassword = '{resetPW}' where userID = {rUserID}"
        # 执行 SQL 并提交
        execute_sql_and_commit(conn, cur, sql)
        # 更新信息，表示密码已重置
        rInfo += "密码已重置为: 1234 / "

    # 如果 rOption2 为真
    if rOption2:
        # 如果 rUserType 有值
        if rUserType:
            # 构建 SQL 更新语句，将用户类型更改为管理员
            sql = f"UPDATE users SET userType = 'admin' where userID = {rUserID}"
            # 更新信息，表示账户类型已更改为管理员
            rInfo += "账户类型已更改为: 管理员 / "
        else:
            # 构建 SQL 更新语句，将用户类型更改为普通用户
            sql = f"UPDATE users SET userType = 'user' where userID = {rUserID}"
            # 更新信息，表示账户类型已更改为用户
            rInfo += "账户类型已更改为: 用户 / "
        # 执行 SQL 并提交
        execute_sql_and_commit(conn, cur, sql)

    # 显示操作结果
    st.success(f"**{rInfo[:-3]}**")


@st.fragment
def task_input():
    st.markdown("### <font face='微软雅黑' color=red><center>工作量录入</center></font>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    expanded_group = ['值班', '清理保洁', '行政管理']
    col1.markdown(f"#### 当前用户: {st.session_state.userCName}")
    with col1:
        flag_auto_task = sac.switch("自动带入默认工作", value=True, align="start", on_label="On")
    task_date = col2.date_input('工作时间', value=datetime.date.today(), max_value="today")
    confirm_btn_input = st.button("确认添加")
    ttl_score = 0
    sql = f"SELECT clerk_work, task_score, task_group from clerk_work where clerk_id = {st.session_state.userID} and task_date = '{task_date}'"
    result = execute_sql(cur, sql)
    if result:
        st.markdown("##### 已输入工作量:\n\n")
        for row in result:
            st.write(f':violet[工作类型:] {row[2]} :orange[内容:] {row[0]} :green[分值:] {row[1]}')
            ttl_score += row[1]
        st.markdown(f':red[总分:] {ttl_score}')
    else:
        st.markdown(f'###### :red[无任何记录]')
    sql = "SELECT DISTINCT(task_group) from bjs_pa"
    rows = execute_sql(cur, sql)
    for row in rows:
        if row[0] in expanded_group and flag_auto_task:
            flag_expanded = True
        else:
            flag_expanded = False
        with st.expander(f"# :green[{row[0]}]", expanded=flag_expanded):
            sql = f"SELECT ID, pa_content, pa_score, pa_group, multi_score, min_days, default_task from bjs_pa where task_group = '{row[0]}' order by pa_num"
            rows2 = execute_sql(cur, sql)
            for row2 in rows2:
                if row2[6] == st.session_state.clerkType and flag_auto_task:
                    auto_task = True
                else:
                    auto_task = False
                if row2[5] > 0:
                    st.checkbox(f":red[{row2[1]} 分值:{row2[2]}]", value=auto_task, key=f"task_work_{row2[0]}")
                else:
                    st.checkbox(f"{row2[1]} 分值:{row2[2]}", value=auto_task, key=f"task_work_{row2[0]}")
                if row2[4] == 1:
                    st.slider(f"倍数", min_value=1, max_value=10, value=1, step=1, key=f"task_multi_{row2[0]}")
    if confirm_btn_input:
        for key in st.session_state.keys():
            if key.startswith("task_work_") and st.session_state[key]:
                #st.write(key, st.session_state[key])
                task_id = key[key.rfind("_") + 1:]
                sql = f"SELECT pa_content, pa_score, task_group from bjs_pa where ID = {task_id}"
                task_result = execute_sql(cur, sql)
                task_content, task_score, task_group = task_result[0]
                if f'task_multi_{task_id}' in st.session_state.keys():
                    task_score *= st.session_state[f'task_multi_{task_id}']
                    #st.write(f"倍数: {st.session_state[f'task_multi_{task_id}']}")
                sql = f"SELECT ID from clerk_work where task_date = '{task_date}' and clerk_id = {st.session_state.userID} and clerk_work = '{task_content}' and task_group = '{task_group}'"
                if not execute_sql(cur, sql):
                    sql = f"INSERT INTO clerk_work (task_date, clerk_id, clerk_cname, clerk_work, task_score, task_group) VALUES ('{task_date}', {st.session_state.userID}, '{st.session_state.userCName}', '{task_content}', {task_score}, '{task_group}')"
                    execute_sql_and_commit(conn, cur, sql)
                    st.toast(f"工作量: [{task_content}] 分值: [{task_score}] 添加成功！")
                else:
                    st.warning(f"工作量: [{task_content}] 已存在！")


def query_task():
    st.markdown("### <font face='微软雅黑' color=red><center>工作量查询及导出</center></font>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    if st.session_state.userType == 'admin':
        userID, userCName = [], []
        sql = "SELECT userID, userCName from users where clerk_pa = 1 order by ID"
        rows = execute_sql(cur, sql)
        for row in rows:
            userID.append(row[0])
            userCName.append(row[1])
        query_userCName = col1.selectbox("请选择查询用户", userCName)
        query_userID = userID[userCName.index(query_userCName)]
    elif st.session_state.userType == 'user':
        col1.markdown(f"#### 当前用户: {st.session_state.userCName}")
        query_userCName = st.session_state.userCName
        query_userID = st.session_state.userID
    query_date_start = col2.date_input('查询开始时间', value=datetime.date.today())
    query_date_end = col3.date_input('查询结束时间', value=datetime.date.today())
    confirm_btn_output = col1.button("导出为Word文件")
    with col2:
        flag_approved = sac.switch("仅限已核定工作", value=False, on_label="On")
    with col3:
        flag_combine = sac.switch("合并统计", value=False, on_label="On")
    if flag_combine:
        sql_task = f"SELECT clerk_work, AVG(task_score) AS avg_task_score, task_group, count(clerk_work) FROM clerk_work WHERE task_approved >= {flag_approved} and task_date >= '{query_date_start}' AND task_date <= '{query_date_end}' AND clerk_id = {query_userID} GROUP BY clerk_work, task_group ORDER BY task_group"
        affix_info, color_field, score_num = "(合并统计)", "单项分值", 1
    else:
        sql_task = f"SELECT task_date, clerk_work, task_score, task_group, task_approved from clerk_work where task_approved >= {flag_approved} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' and clerk_id = {query_userID} order by task_date, task_group, ID, clerk_work"
        affix_info, color_field, score_num = "", "分值", 2
    rows = execute_sql(cur, sql_task)
    if rows:
        ttl_score = 0
        for row in rows:
            if flag_combine:
                ttl_score = ttl_score + row[1] * row[3]
            else:
                ttl_score += row[2]
        ttl_score = int(ttl_score)
        df = pd.DataFrame(rows, dtype=str)
        if flag_combine:
            df.columns = ["工作", "单项分值", "工作组别", "工作项数"]
            for index, value in enumerate(rows):
                df.loc[index, "单项分值"] = int(float(df["单项分值"][index]))
        else:
            df.columns = ["日期", "工作", "分值", "工作组别", "核定状态"]
            for index, value in enumerate(rows):
                df.loc[index, "分值"] = int(df["分值"][index])
                df.loc[index, "核定状态"] = "已核定" if int(df["核定状态"][index]) == 1 else "未核定"
        st.dataframe(df.style.apply(highlight_max, subset=[color_field]))
        #st.dataframe(df)
        st.markdown(f':green[共计:] {len(rows)}项工作{affix_info} :red[总分:] {ttl_score}')
    else:
        st.info(f":red[没有查询到符合条件的记录]")
    if confirm_btn_output:
        headerFS = 16
        contentFS = 12
        quesDOC = Document()
        quesDOC.styles["Normal"].font.name = "Microsoft YaHei"
        quesDOC.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        pHeader = quesDOC.add_paragraph()
        pHeader.alignment = WD_ALIGN_PARAGRAPH.CENTER
        textHeader = pHeader.add_run(f"{st.session_state.userCName} {query_date_start} 至 {query_date_end} 工作量{affix_info}记录", 0)
        textHeader.font.size = Pt(headerFS)
        textHeader.font.bold = True
        rows = execute_sql(cur, sql_task)
        if rows:
            i = 1
            for row in rows:
                pContent = quesDOC.add_paragraph()
                if flag_combine:
                    textContent = pContent.add_run(f"第{i}项 - 工作类型: {row[2]} 内容: {row[0]} 单项分值: {int(row[1])} 项数: {row[3]}次")
                else:
                    if row[4] == 1:
                        approved_text = "已核定"
                    else:
                        approved_text = "未核定"
                    textContent = pContent.add_run(f"第{i}项 - 日期: {row[0]} 工作类型: {row[3]} 内容: {row[1]} 分值: {row[2]} 状态: {approved_text}")
                textContent.font.size = Pt(contentFS)
                textContent.font.bold = False
                if row[score_num] < 0:
                    textContent.font.color.rgb = RGBColor(139, 0, 0)
                else:
                    textContent.font.color.rgb = RGBColor(0, 0, 0)
                i += 1
            sql = f"SELECT clerk_work, task_score, task_group from clerk_work where task_date >= '{query_date_start}' and task_date <= '{query_date_end}' and clerk_id = {query_userID} and task_score < 0 order by task_date, task_group, ID, clerk_work"
            deduct_result = execute_sql(cur, sql)
            if deduct_result:
                pContent = quesDOC.add_paragraph()
                textContent = pContent.add_run(f"其中共{len(deduct_result)}个减分项(未合并)")
                textContent.font.size = Pt(contentFS + 2)
                textContent.font.bold = True
                textContent.font.color.rgb = RGBColor(139, 0, 0)
                j = 1
                for deduct_row in deduct_result:
                    pContent = quesDOC.add_paragraph()
                    textContent = pContent.add_run(f"第{j}项 - 工作类型: {deduct_row[2]} 内容: {deduct_row[0]} 分值: {deduct_row[1]}")
                    textContent.font.size = Pt(contentFS)
                    textContent.font.bold = True
                    textContent.font.color.rgb = RGBColor(0, 0, 0)
                    j += 1
            for j in range(1):
                pContent = quesDOC.add_paragraph()
            pContent = quesDOC.add_paragraph()
            textContent = pContent.add_run(f"共计完成{i - 1}项工作{affix_info} 总分: {ttl_score}")
            textContent.font.size = Pt(contentFS + 2)
            textContent.font.bold = True
            textContent.font.color.rgb = RGBColor(155, 17, 30)
            for j in range(1):
                pContent = quesDOC.add_paragraph()
            pContent = quesDOC.add_paragraph()
            textContent = pContent.add_run("核定签字: ________")
            add_page_number(quesDOC.sections[0].footer.paragraphs[0].add_run())
            quesDOC.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            outputFile = f"./user_pa/{st.session_state.userCName}-{query_date_start}-{query_date_end}_{time.strftime('%Y%m%d%H%M%S', time.localtime(int(time.time())))}.docx"
            if os.path.exists(outputFile):
                os.remove(outputFile)
            quesDOC.save(outputFile)
            if os.path.exists(outputFile):
                if os.path.exists(outputFile):
                    with open(outputFile, "rb") as file:
                        content = file.read()
                    file.close()
                    buttonDL = st.download_button("点击下载", content, file_name=outputFile[outputFile.rfind("/") + 1:], icon=":material/download:", type="secondary")
                    st.success(f":green[成功导出至程序目录user_pa下 {outputFile[outputFile.rfind('/') + 1:]}]")
                    if buttonDL:
                        st.toast("文件已下载至你的默认目录")
            else:
                st.error(f":red[文件导出失败]")
        else:
            st.info(f":red[没有查询到符合条件的记录]")

def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(qn(name), value)


def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def manual_input():
    items = []
    st.markdown("### <font face='微软雅黑' color=red><center>工作量手工录入</center></font>", unsafe_allow_html=True)
    st.markdown(f"#### 当前用户: {st.session_state.userCName}")
    sql = "SELECT DISTINCT(task_group) from bjs_pa"
    rows = execute_sql(cur, sql)
    for row in rows:
        items.append(row[0])
    col1, col2, col3 = st.columns(3)
    task_date = col1.date_input('工作时间', value=datetime.date.today(), max_value="today")
    task_group = col2.selectbox('工作组别', items, index=None, accept_new_options=True)
    task_score = col3.slider("单项分值", min_value=5, max_value=300, value=10, step=5)
    opt1, opt2, opt3 = st.columns(3)
    if st.session_state.userType == 'admin':
        with opt1:
            flag_add_pa = sac.switch("加入固定列表", value=False, align="start", on_label="On")
        with opt2:
            flag_multi_score = sac.switch("多倍计算", value=False, align="start", on_label="On")
    else:
        flag_add_pa, flag_multi_score = False, False
    task_content = st.text_area("工作内容", height=100)
    confirm_btn_manual = st.button("确认添加")
    if task_group and task_content and confirm_btn_manual:
        sql = f"SELECT ID from clerk_work where task_date = '{task_date}' and clerk_id = {st.session_state.userID} and clerk_work = '{task_content}' and task_group = '{task_group}'"
        if not execute_sql(cur, sql):
            sql = f"INSERT INTO clerk_work (task_date, clerk_id, clerk_cname, clerk_work, task_score, task_group) VALUES ('{task_date}', {st.session_state.userID}, '{st.session_state.userCName}', '{task_content}', {task_score}, '{task_group}')"
            execute_sql_and_commit(conn, cur, sql)
            st.toast(f"工作量: [{task_content}] 添加成功！")
        else:
            st.warning(f"工作量: [{task_content}] 已存在！")
        if flag_add_pa:
            sql = f"SELECT ID from bjs_pa where pa_content = '{task_content}' and task_group = '{task_group}' and pa_score = {task_score}"
            if not execute_sql(cur, sql):
                sql = f"INSERT INTO bjs_pa (pa_content, pa_score, pa_group, task_group, multi_score) VALUES ('{task_content}', {task_score}, '全员', '{task_group}', {int(flag_multi_score)})"
                execute_sql_and_commit(conn, cur, sql)
                reset_table_num(True)
                st.toast(f"工作量: [{task_content}] 添加至列表成功！")
            else:
                st.warning(f"工作量: [{task_content}] 在列表中已存在！")
    elif not task_group:
        st.warning(f"请选择工作组！")
    elif not task_content:
        st.warning(f"请输入工作内容！")

def reset_table_num(flag_force=False):
    if not flag_force:
        confirm_btn_reset = st.button("确认重置")
    else:
        confirm_btn_reset = True
    if confirm_btn_reset:
        for modify_table in ['bjs_pa', 'bjs_pa_deduct']:
            i, sql = 1, ''
            if modify_table == 'bjs_pa':
                sql = f"SELECT ID from {modify_table} order by task_group, ID, pa_num"
            elif modify_table == 'bjs_pa_deduct':
                sql = f"SELECT ID from {modify_table} order by ID"
            if sql:
                rows = execute_sql(cur, sql)
                for row in rows:
                    sql = f"UPDATE {modify_table} SET pa_num = {i} where ID = {row[0]}"
                    execute_sql_and_commit(conn, cur, sql)
                    i += 2
        if not flag_force:
            st.success(f"{modify_table} ID重置成功")


#@st.fragment
def task_modify():
    st.markdown("### <font face='微软雅黑' color=red><center>记录修改</center></font>", unsafe_allow_html=True)
    col1, col2, col3, col4 = st.columns(4)
    if st.session_state.userType == 'admin':
        userID, userCName = [], []
        sql = "SELECT userID, userCName from users where clerk_pa = 1 order by ID"
        rows = execute_sql(cur, sql)
        for row in rows:
            userID.append(row[0])
            userCName.append(row[1])
        query_userCName = col1.selectbox("请选择查询用户", userCName)
        query_userID = userID[userCName.index(query_userCName)]
    else:
        col1.markdown(f"#### 当前用户: {st.session_state.userCName}")
        query_userID = st.session_state.userID
        query_userCName = st.session_state.userCName
    query_date_start = col2.date_input('查询开始时间', value=datetime.date.today(), max_value="today")
    query_date_end = col3.date_input('查询结束时间', value=datetime.date.today(), max_value="today")
    user_task_id_pack = []
    sql = f"SELECT clerk_work, task_score, task_group, ID from clerk_work where clerk_id = {query_userID} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}'"
    result = execute_sql(cur, sql)
    for row in result:
        user_task_id_pack.append(row[3])
    task_modify_id = col4.selectbox("请选择任务ID", user_task_id_pack, index=None)
    form = st.columns(3)
    confirm_btn_delete = form[0].button("删除", type="primary")
    display_are = st.empty()
    with display_are.container(border=True):
        if result:
            ttl_score = 0
            st.markdown("##### 已输入工作量:\n\n")
            for row in result:
                st.write(f'ID:{row[3]} :violet[工作类型:] {row[2]} :orange[内容:] {row[0]} :green[分值:] {row[1]}')
                ttl_score += row[1]
            st.markdown(f':red[总分:] {ttl_score}')
        else:
            st.markdown(f'###### :red[无任何记录]')
    if task_modify_id:
        if confirm_btn_delete:
            form[1].button("确认删除", type="secondary", on_click=delete_task, args=(task_modify_id, query_userID,))
        if st.session_state.userType == 'admin':
            display_are.empty()
            modify_task(task_modify_id, query_userID)
    else:
        st.info('请选择要处理的记录ID')


def delete_task(task_modify_id, query_userID):
    if st.session_state.userType == 'admin':
        sql = f"DELETE FROM clerk_work where ID = {task_modify_id} and clerk_id = {query_userID}"
    else:
        sql = f"DELETE FROM clerk_work where ID = {task_modify_id} and clerk_id = {query_userID} and task_approved = 0"
    execute_sql_and_commit(conn, cur, sql)
    sql = f"SELECT ID FROM clerk_work where ID = {task_modify_id} and clerk_id = {query_userID}"
    if not execute_sql(cur, sql):
        st.toast(f"ID:{task_modify_id} 删除成功!")
    else:
        st.toast(f"ID:{task_modify_id} 删除失败! 被核定的记录无法删除, 请联系管理员!")


def modify_task(task_modify_id, query_userID):
    sql = f"SELECT clerk_work, task_score, task_group from clerk_work where ID = {task_modify_id} and clerk_id = {query_userID}"
    modify_pack = execute_sql(cur, sql)[0]
    form = st.columns(3)
    modify_content = form[0].text_area("请输入修改后的内容", value=modify_pack[0], height=100)
    modify_score = form[1].number_input("请输入修改后的分数", min_value=MAXDEDUCTSCORE, max_value=1000, value=modify_pack[1], step=1, placeholder="最大1000")
    sql = f"UPDATE clerk_work SET clerk_work = '{modify_content}', task_score = {modify_score} where ID = {task_modify_id} and clerk_id = {query_userID}"
    execute_sql_and_commit(conn, cur, sql)
    sql = f"SELECT ID from clerk_work where clerk_work = '{modify_content}' and task_score = {modify_score} and ID = {task_modify_id} and clerk_id = {query_userID}"
    if execute_sql(cur, sql):
        pass
        #st.toast(f"ID:{task_modify_id} 修改成功!")
    else:
        st.toast(f"ID:{task_modify_id} 修改失败! 被核定的记录无法修改, 请联系管理员!")

@st.fragment
def check_data():
    st.markdown("### <font face='微软雅黑' color=red><center>数据检查与核定</center></font>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    userID, userCName = [], []
    sql = "SELECT userID, userCName from users where clerk_pa = 1 order by ID"
    rows = execute_sql(cur, sql)
    for row in rows:
        userID.append(row[0])
        userCName.append(row[1])
    query_date_start = col1.date_input('查询开始时间', value=datetime.date.today())
    query_date_end = col2.date_input('查询结束时间', value=datetime.date.today())
    dur_time = query_date_end - query_date_start
    st.markdown(f'##### 统计周期: {dur_time.days}天')
    confirm_btn_check = col1.button("检查")
    confirm_btn_approv = col2.button("核定")
    if confirm_btn_check:
        for index, value in enumerate(userID):
            sql = "SELECT pa_content, min_days from bjs_pa where min_days > 0 order by min_days DESC"
            rows = execute_sql(cur, sql)
            for row in rows:
                sql = f"SELECT count(ID) from clerk_work where clerk_work = '{row[0]}' and clerk_id = {value} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}'"
                task_count = execute_sql(cur, sql)[0][0]
                if task_count > 1 and task_count > dur_time.days / row[1]:
                    st.warning(f"用户: {userCName[index]} 工作: [{row[0]}] 应该 1次/{row[1]}天, 实际: {task_count}次 已超量, 请检查记录！")
    else:
        task_pack = []
        sql = f"SELECT ID, clerk_cname, task_date, clerk_work, task_score, task_group from clerk_work where task_approved = 0 and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' order by task_date, task_group, clerk_work, clerk_cname, task_score"
        result = execute_sql(cur, sql)
        if result:
            for row in result:
                task_pack.append(f'日期:{row[2]} 用户:{row[1]} 工作类型:{row[5]} 内容:{row[3]} 分值:{row[4]} ID:{row[0]}')
            approve_pack = sac.transfer(items=task_pack, label='工作量核定', titles=['项目'], reload=True, align='center', search=True, pagination=True, use_container_width=True)
            if confirm_btn_approv and approve_pack:
                for each in approve_pack:
                    approve_id = each[each.find('ID:') + 3:].strip()
                    sql = f"UPDATE clerk_work SET task_approved = 1 where ID = {approve_id}"
                    execute_sql_and_commit(conn, cur, sql)
                    st.success(f"**{each} 工作量已核定**")
        else:
            st.markdown(f'###### :red[无任何记录]')

def resetPassword():
    # 显示副标题和分隔线
    st.subheader(":orange[密码重置]", divider="red")

    # 检查是否需要重置用户信息
    if st.session_state.userPwRechecked:
        # 显示重置用户信息提示
        st.write(":red[**重置用户信息**]")

        # 创建三列布局
        rCol1, rCol2 = st.columns(2)

        # 获取用户编码
        rUserName = rCol1.number_input("用户编码", min_value=1, max_value=20, value=1)

        # 执行SQL查询用户信息
        sql = f"SELECT userCName, userType from users where userID = {rUserName}"
        rows = execute_sql(cur, sql)

        # 检查是否查询到用户信息
        if rows:
            # 显示用户姓名
            rCol2.write(f"用户姓名: **{rows[0][0]}**")

            # 创建重置按钮
            btnResetUserPW = st.button("重置", type="primary")

            if btnResetUserPW:
                st.button("确认重置", type="secondary", on_click=actionResetUserPW, args=(rUserName,))
                st.session_state.userPwRechecked = False
        # 如果未查询到用户信息，显示错误
        else:
            st.error("用户不存在")
    else:
        vUserPW = st.text_input("请输入密码", max_chars=8, placeholder="请输入管理员密码, 以验证身份", type="password", autocomplete="off")

        # 检查是否输入了密码
        if vUserPW:
            # 验证密码
            if verifyUserPW(st.session_state.userID, vUserPW)[0]:
                st.session_state.userPwRechecked = True
                st.rerun()
            # 如果密码错误，显示错误提示
            else:
                st.session_state.userPwRechecked = False
                st.error("密码错误, 请重新输入")

def actionResetUserPW(rUserName):
    rInfo = ""

    # 获取用户加密密钥
    resetPW = getUserEDKeys("1234", "enc")
    # 构建 SQL 更新语句
    sql = f"UPDATE users SET userPassword = '{resetPW}' where userID = {rUserName}"
    # 执行 SQL 并提交
    execute_sql_and_commit(conn, cur, sql)
    # 更新信息，表示密码已重置
    rInfo += "密码已重置为: 1234"

    st.success(f"**{rInfo}**")


def deduction_input():
    st.markdown("### <font face='微软雅黑' color=red><center>减分项录入</center></font>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    userID, userCName, pa_deduct, pa_deduct_score = [], [], [], []
    sql = "SELECT userID, userCName from users where clerk_pa = 1 order by ID"
    rows = execute_sql(cur, sql)
    for row in rows:
        userID.append(row[0])
        userCName.append(row[1])
    deduct_userCName = col1.selectbox("请选择用户", userCName)
    deduct_userID = userID[userCName.index(deduct_userCName)]
    deduct_date = col2.date_input("请选择日期", datetime.date.today(), max_value="today")
    sql = "SELECT pa_content, pa_score from bjs_pa_deduct order by ID"
    rows = execute_sql(cur, sql)
    for row in rows:
        pa_deduct.append(row[0])
        pa_deduct_score.append(row[1])
    task_deduct = col1.selectbox("扣分项", pa_deduct, index=None)
    if task_deduct:
        task_score = pa_deduct_score[pa_deduct.index(task_deduct)]
    else:
        task_score = -1
    deduct_score = col2.number_input("扣分", min_value=MAXDEDUCTSCORE, max_value=-1, value=task_score, step=1, placeholder=f"最小值{MAXDEDUCTSCORE}, 最大值-1")
    deduct_content = col1.text_area("自定义扣分项内容", value=task_deduct, placeholder="可选择固定扣分项后修改", height=100)
    confirm_btn_add = st.button("确认添加")
    if confirm_btn_add:
        #st.write(deduct_content, deduct_score, deduct_userID, deduct_userCName, deduct_date)
        if deduct_content:
            sql = f"SELECT ID from clerk_work where task_date = '{deduct_date}' and clerk_work = '{deduct_content}' and clerk_id = {deduct_userID}"
            if not execute_sql(cur, sql):
                sql = f"INSERT INTO clerk_work (task_date, clerk_id, clerk_cname, clerk_work, task_score, task_group, task_approved) VALUES ('{deduct_date}', {deduct_userID}, '{deduct_userCName}', '{deduct_content}', {deduct_score}, '扣分', 1)"
                execute_sql_and_commit(conn, cur, sql)
                st.success(f"{deduct_userCName} 扣分项添加成功")
            else:
                st.error(f"{deduct_userCName} {deduct_date} {deduct_content} 扣分项已存在")
            sql = f"SELECT ID from bjs_pa_deduct where pa_content = '{deduct_content}' and pa_score = {deduct_score}"
            if not execute_sql(cur, sql):
                sql = f"INSERT INTO bjs_pa_deduct(pa_content, pa_score) VALUES ('{deduct_content}', {deduct_score})"
                execute_sql_and_commit(conn, cur, sql)
                st.success(f"{deduct_content} 扣分项已添加至固定列表")
                reset_table_num(True)
        else:
            st.error("请输入扣分项内容")


def highlight_max(x, forecolor='black', backcolor="#D61919"):
    is_max = x == x.max()

    return [f'color: {forecolor}; background-color: {backcolor}' if v else '' for v in is_max]


def gen_chart():
    st.markdown("### <font face='微软雅黑' color=red><center>趋势图</center></font>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    tab1, tab2 = st.tabs(["📈 图表", "🗃 数据"])
    if st.session_state.userType == 'admin':
        userID, userCName = [], []
        sql = "SELECT userID, userCName from users where clerk_pa = 1 order by ID"
        rows = execute_sql(cur, sql)
        for row in rows:
            userID.append(row[0])
            userCName.append(row[1])
        query_userCName = col1.selectbox("请选择查询用户", userCName)
        query_userID = userID[userCName.index(query_userCName)]
    elif st.session_state.userType == 'user':
        col1.markdown(f"#### 当前用户: {st.session_state.userCName}")
        query_userCName = st.session_state.userCName
        query_userID = st.session_state.userID
    query_date_start = col2.date_input('查询开始时间', value=datetime.date.today())
    query_date_end = col3.date_input('查询结束时间', value=datetime.date.today())
    if st.session_state.userType == 'admin':
        with col1:
            flag_all_user = sac.switch("查询所有用户", value=False, align="start", on_label="On")
    else:
        flag_all_user = False
    with col2:
        flag_approved = sac.switch("仅限已核定工作", value=False, on_label="On")
    if not flag_all_user:
        userID = [query_userID]
        userCName = [query_userCName]
    with col3:
        #dur_time = query_date_end - query_date_start
        chart_type_pack = ['折线图', '旭日图', '矩阵树图', '饼图']
        if len(userID) == 1:
            chart_type_pack = chart_type_pack + ['热度图', '柱状图(分组)', '柱状图(堆叠)', '漏斗图']
        chart_type = st.selectbox("图表类型", chart_type_pack, index=0, placeholder="多人查询只有折线图、旭日图、矩阵树图和饼图")
    min_value, max_value = 1000, 0
    raws_data, df = [], []
    charArea = tab1.empty()
    if chart_type == '折线图':
        with tab1:
            # 双Y轴折线图
            with charArea.container(border=True):
                fig = go.Figure()
                for index, value in enumerate(userID):
                    hot_value, hot_date, temp_value_pack = [], [], []
                    sql = f"SELECT task_date, sum(task_score) from clerk_work where task_approved >= {flag_approved} and clerk_id = {value} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' GROUP BY task_date order by task_date"
                    result = execute_sql(cur, sql)
                    for each in result:
                        hot_date.append(each[0])
                        hot_value.append(int(each[1]))
                        raws_data.append([userCName[index], each[0], int(each[1])])
                    temp_value_pack = hot_value
                    if temp_value_pack:
                        temp_value_pack.sort()
                        if temp_value_pack[0] < min_value:
                            min_value = temp_value_pack[0]
                        if temp_value_pack[-1] > max_value:
                            max_value = temp_value_pack[-1]
                        #st.write(min_value, max_value, hot_value)
                        if temp_value_pack[-1] < max_value / 2:
                            yax = 'y2'
                        else:
                            yax = 'y'
                        fig.add_trace(
                            go.Scatter(name=f"{userCName[index]}",
                                        x=hot_date,
                                        y=hot_value,
                                        mode="markers+lines+text",
                                        text=[
                                            format(round(x, 2), ",")
                                            for x in hot_value
                                        ],
                                        yaxis=yax,
                                        textposition="top center"))
                if raws_data:
                    fig.update_layout(
                        title="工作量",
                        xaxis=dict(title="日期"),
                        yaxis=dict(title="主轴",
                                    rangemode="normal"),
                        template="simple_white",
                        yaxis2=dict(
                            title="",
                            overlaying='y',
                            side='right'))
                    st.plotly_chart(fig)
                    df = pd.DataFrame(raws_data, columns=["姓名", "日期", "合计分值"])
                else:
                    st.info("没有查询到符合条件的记录")
    elif chart_type.startswith("柱状图"):
        if "分组" in chart_type:
            bar_type = "group"
        elif "堆叠" in chart_type:
            bar_type = "stack"
        else:
            bar_type = "relative"
        with tab1:
            with charArea.container(border=True):
                for index, value in enumerate(userID):
                    sql = f"SELECT task_date, task_group, sum(task_score) from clerk_work where task_approved >= {flag_approved} and clerk_id = {value} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' GROUP BY task_date, task_group order by task_date"
                    result = execute_sql(cur, sql)
                    for each in result:
                        raws_data.append([userCName[index], each[0], each[1], int(each[2])])
                if raws_data:
                    df = pd.DataFrame(raws_data, columns=["姓名", "日期", "工作组别", "合计分值"])
                    # 使用 Plotly Express 生成分组柱状图
                    fig = px.bar(
                        df,
                        x="日期",
                        y="合计分值",
                        color="工作组别",
                        text="合计分值",
                        title="按日期和工作组别统计",
                        labels={"合计分值": "总分", "日期": "工作日期", "工作组别": "任务组"},
                        barmode=bar_type
                    )
                    # 调整样式
                    fig.update_traces(textposition='outside')
                    fig.update_layout(
                        xaxis_tickangle=-45,
                        template="simple_white"
                    )
                    st.plotly_chart(fig)
                else:
                    st.info("没有查询到符合条件的记录")
    elif chart_type == "热度图":
        with tab1:
            with charArea.container(border=True):
                for index, value in enumerate(userID):
                    sql = f"SELECT task_date, task_group, sum(task_score) from clerk_work where task_approved >= {flag_approved} and clerk_id = {value} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' GROUP BY task_date, task_group order by task_date"
                    result = execute_sql(cur, sql)
                    for each in result:
                        raws_data.append([userCName[index], each[0], each[1], int(each[2])])
                if raws_data:
                    df = pd.DataFrame(raws_data, columns=["姓名", "日期", "工作组别", "合计分值"])
                    # 构建透视表
                    heatmap_data = df.pivot_table(index="工作组别", columns="日期", values="合计分值", aggfunc="sum")
                    # 生成热图
                    fig = px.imshow(
                        heatmap_data,
                        text_auto=True,
                        color_continuous_scale='Viridis',
                        title="工作量",
                        labels={"x": "日期", "y": "工作组别", "color": "总分"}
                    )
                    st.plotly_chart(fig)
                else:
                    st.info("没有查询到符合条件的记录")
    elif chart_type == "漏斗图":
        with tab1:
            with charArea.container(border=True):
                for index, value in enumerate(userID):
                    sql = f"SELECT task_date, task_group, sum(task_score) from clerk_work where task_approved >= {flag_approved} and clerk_id = {value} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' GROUP BY task_date, task_group order by sum(task_score) DESC"
                    result = execute_sql(cur, sql)
                    for each in result:
                        raws_data.append([userCName[index], each[0], each[1], int(each[2])])
                if raws_data:
                    df = pd.DataFrame(raws_data, columns=["姓名", "日期", "工作组别", "合计分值"])
                    # 按工作组别统计总分并按降序排序
                    funnel_data = df.groupby("工作组别")["合计分值"].sum().reset_index()
                    funnel_data = funnel_data.sort_values(by="合计分值", ascending=False)
                    # 生成漏斗图
                    fig = px.funnel(
                        funnel_data,
                        x="合计分值",
                        y="工作组别",
                        title="工作量(日期合并)",
                        labels={"合计分值": "总分", "工作组别": "任务组别"},
                        color_discrete_sequence=px.colors.qualitative.Prism
                    )
                    st.plotly_chart(fig)
                else:
                    st.info("没有查询到符合条件的记录")
    elif chart_type == "饼图":
        with tab1:
            with charArea.container(border=True):
                for index, value in enumerate(userID):
                    sql = f"SELECT task_group, sum(task_score) from clerk_work where task_approved >= {flag_approved} and clerk_id = {value} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' GROUP BY task_group order by sum(task_score) DESC"
                    result = execute_sql(cur, sql)
                    for each in result:
                        raws_data.append([userCName[index], each[0], int(each[1])])
                if raws_data:
                    df = pd.DataFrame(raws_data, columns=["姓名", "工作组别", "合计分值"])
                    if len(userID) > 1:
                        pie_data = df.groupby("工作组别")["合计分值"].sum().reset_index()
                    else:
                        pie_data = df.copy()
                    # 计算总和
                    total = pie_data['合计分值'].sum()
                    # 添加百分比列
                    pie_data['百分比'] = (pie_data['合计分值'] / total) * 100
                    # 保留所有原始条目，不进行合并
                    final_data = pie_data.copy()
                    fig = px.pie(
                        final_data,
                        names="工作组别",
                        values="合计分值",
                        title="工作量(日期合并)",
                        hole=0.2,
                        hover_data=["合计分值"],
                        labels={"合计分值": "总分", "工作组别": "任务组别"},
                        color_discrete_sequence=px.colors.qualitative.Prism
                    )
                    fig.update_traces(textposition='outside', textinfo='percent+label')
                    fig.update_layout(showlegend=False)
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("没有查询到符合条件的记录")
    elif chart_type == "旭日图":
        with tab1:
            with charArea.container(border=True):
                for index, value in enumerate(userID):
                    # 查询每个用户的任务分值按工作组别汇总
                    sql = f"SELECT task_group, sum(task_score) from clerk_work where task_approved >= {flag_approved} and clerk_id = {value} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' GROUP BY task_group order by sum(task_score) DESC"
                    result = execute_sql(cur, sql)
                    for each in result:
                        raws_data.append([userCName[index], each[0], int(each[1])])
                if raws_data:
                    # 构造 DataFrame
                    df = pd.DataFrame(raws_data, columns=["姓名", "工作组别", "合计分值"])
                    if len(userID) > 1:
                        sunburst_data = df.groupby(["工作组别", "姓名"], as_index=False)["合计分值"].sum()
                    else:
                        sunburst_data = df.copy()
                    # 绘制旭日图
                    fig = px.sunburst(
                        sunburst_data,
                        path=['工作组别', '姓名'],
                        values='合计分值',
                        color='合计分值',
                        hover_data=['合计分值'],
                        color_continuous_scale='Plasma',
                        title="工作量（工作组别 → 用户）",
                    )
                    fig.update_layout(margin=dict(t=50, l=0, r=0, b=0))
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("没有查询到符合条件的记录")
    elif chart_type == "矩阵树图":
        with tab1:
            with charArea.container(border=True):
                for index, value in enumerate(userID):
                    # 查询每个用户的任务分值按工作组别汇总
                    sql = f"SELECT task_group, sum(task_score) from clerk_work where task_approved >= {flag_approved} and clerk_id = {value} and task_date >= '{query_date_start}' and task_date <= '{query_date_end}' GROUP BY task_group order by sum(task_score) DESC"
                    result = execute_sql(cur, sql)
                    for each in result:
                        raws_data.append([userCName[index], each[0], int(each[1])])
                if raws_data:
                    df = pd.DataFrame(raws_data, columns=["姓名", "工作组别", "合计分值"])
                    if len(userID) > 1:
                        treemap_data = df.groupby(["工作组别", "姓名"], as_index=False)["合计分值"].sum()
                    else:
                        treemap_data = df.copy()
                    fig = px.treemap(
                        treemap_data,
                        path=['工作组别', '姓名'],
                        values='合计分值',
                        color='合计分值',
                        color_continuous_scale='Plasma',
                        title="工作量（工作组别 → 用户）",
                        hover_data={'合计分值': True}
                    )
                    fig.update_layout(margin=dict(t=50, l=0, r=0, b=0))
                    st.plotly_chart(fig, use_container_width=True)
                else:
                    st.info("没有查询到符合条件的记录")
    if raws_data:
        tab2.write(df)
    else:
        tab2.info("没有查询到符合条件的记录")

@st.fragment
def displayBigTime():
    components.html(open("./MyComponentsScript/Clock-Big.txt", "r", encoding="utf-8").read(), height=140)


global APPNAME, MAXDEDUCTSCORE
APPNAME = "北京站绩效考核系统KPI-PA"
MAXDEDUCTSCORE = -20
conn = get_connection()
cur = conn.cursor()
st.logo("./Images/logos/bjs-pa-logo2.png", icon_image="./Images/logos/bjs-pa-logo.png", size="large")

selected = None

if "logged_in" not in st.session_state:
    st.set_page_config(layout='wide')
    st.session_state.logged_in = False
    login()

if st.session_state.logged_in:
    with st.sidebar:
        #displaySmallTime()
        #displaySmallClock()
        if st.session_state.userType == "admin":
            selected = sac.menu([
                sac.MenuItem('主页', icon='house'),
                sac.MenuItem('功能', icon='grid-3x3-gap', children=[
                    sac.MenuItem('工作量录入', icon='list-task'),
                    sac.MenuItem('工作量手工录入', icon='journal-plus'),
                    sac.MenuItem('工作减分项录入', icon='journal-minus'),
                    sac.MenuItem('记录修改', icon='journal-medical'),
                    sac.MenuItem('统计查询及导出', icon='clipboard-data'),
                    sac.MenuItem('趋势图', icon='bar-chart-line'),
                    sac.MenuItem('数据检查与核定', icon='check2-all'),
                    sac.MenuItem("重置数据库ID", icon="bootstrap-reboot"),
                ]),
                sac.MenuItem('账户', icon='person-gear', children=[
                    sac.MenuItem('密码修改', icon='key'),
                    sac.MenuItem('密码重置', icon='person-gear'),
                    sac.MenuItem('登出', icon='box-arrow-right'),
                ]),
                sac.MenuItem('关于', icon='layout-wtf', children=[
                    sac.MenuItem('Changelog', icon='view-list'),
                    sac.MenuItem('Readme', icon='github'),
                    sac.MenuItem('关于...', icon='link-45deg'),
                ]),
            ], open_index=[1])
        elif st.session_state.userType == "user":
            selected = sac.menu([
                sac.MenuItem('主页', icon='house'),
                sac.MenuItem('功能', icon='grid-3x3-gap', children=[
                    sac.MenuItem('工作量录入', icon='list-task'),
                    sac.MenuItem('工作量手工录入', icon='journal-plus'),
                    sac.MenuItem('记录修改', icon='journal-medical'),
                    sac.MenuItem('统计查询及导出', icon='clipboard-data'),
                    sac.MenuItem('趋势图', icon='bar-chart-line'),
                ]),
                sac.MenuItem('账户', icon='person-gear', children=[
                    sac.MenuItem('密码修改', icon='key'),
                    sac.MenuItem('登出', icon='box-arrow-right'),
                ]),
                sac.MenuItem('关于', icon='layout-wtf', children=[
                    sac.MenuItem('Changelog', icon='view-list'),
                    sac.MenuItem('Readme', icon='github'),
                    sac.MenuItem('关于...', icon='link-45deg'),
                ]),
            ], open_index=[1])
        st.divider()
        st.markdown(f'### :green[当前用户:] :orange[{st.session_state.userCName}]')
    if selected == "主页":
        #displayBigTimeCircle()
        displayBigTime()
        displayAppInfo()
        displayVisitCounter()
    elif selected == "工作量录入":
        task_input()
    elif selected == "工作量手工录入":
        manual_input()
    elif selected == "工作减分项录入":
        deduction_input()
    elif selected == "记录修改":
        task_modify()
    elif selected == "统计查询及导出":
        query_task()
    elif selected == "趋势图":
        gen_chart()
    elif selected == "数据检查与核定":
        check_data()
    elif selected == "重置数据库ID":
        reset_table_num()
    elif selected == "密码修改":
        changePassword()
    elif selected == "密码重置":
        resetPassword()
    elif selected == "登出":
        logout()
    elif selected == "Changelog":
        changelog()
    elif selected == "Readme":
        aboutReadme()
    elif selected == "关于...":
        aboutInfo()
